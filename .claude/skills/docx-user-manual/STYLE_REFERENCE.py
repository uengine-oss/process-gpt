"""DOCX user manual style reference.

This compact reference is derived from the polished generator previously kept
under docs/prev/docs/generate_text2sql_user_manual.py. Use it as the visual
baseline for new manual generators: copy the helper structure and adapt only
content, screenshot keys, output paths, and feature-specific wording.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[3]
OUTPUT_PATH = REPO_ROOT / "docs" / "<suite-or-feature>-user-manual.docx"
SCREENSHOT_DIR = REPO_ROOT / "e2e" / "<suite-slug>" / "e2e-results" / "<suite-slug>" / "screenshots"

DOC_VERSION = "v1.0"
DOC_DATE = date.today()
ORG_NAME = "K-Water AIR"
SERVICE_NAME = "<서비스명>"
DOC_TITLE = "<기능명> 사용자 매뉴얼"
FONT_NAME = "Malgun Gothic"
PRIMARY_COLOR = "1F4E79"
CAPTION_COLOR = "666666"
TABLE_HEADER_FILL = "D9EAF7"


SCREENSHOTS = {
    "progress": "<suite>-01-progress.png",
    "result": "<suite>-01-result.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차는 Word에서 필드 업데이트 시 페이지 번호와 함께 표시됩니다."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, bold=True, color=PRIMARY_COLOR)


def add_paragraph(doc: Document, text: str = "", style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        run = paragraph.add_run(f"{idx}.\t{item}")
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, TABLE_HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_image(doc: Document, key: str, caption: str) -> None:
    path = SCREENSHOT_DIR / SCREENSHOTS[key]
    if not path.exists():
        add_paragraph(doc, f"[화면 예시 누락] {path.name}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=CAPTION_COLOR)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SERVICE_NAME)
    set_run_font(run, size=18, bold=True, color=PRIMARY_COLOR)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DOC_TITLE)
    set_run_font(run, size=28, bold=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 버전", DOC_VERSION],
            ["작성일", DOC_DATE.isoformat()],
            ["작성자/조직", ORG_NAME],
            ["문서 구분", "사용자 참조 매뉴얼"],
            ["대상 서비스", SERVICE_NAME],
        ],
        [4.0, 12.0],
    )
    doc.add_page_break()


def add_scenario(
    doc: Document,
    title: str,
    purpose: str,
    preconditions: list[str],
    steps: list[str],
    expected: list[str],
    screenshot_key: str | None = None,
    caption: str | None = None,
) -> None:
    add_heading(doc, title, level=2)
    add_paragraph(doc, f"이럴 때 사용합니다: {purpose}")
    add_paragraph(doc, "사용 전 확인")
    add_bullets(doc, preconditions)
    add_paragraph(doc, "따라 하기")
    add_numbers(doc, steps)
    add_paragraph(doc, "화면에서 확인할 내용")
    add_bullets(doc, expected)
    if screenshot_key and caption:
        add_image(doc, screenshot_key, caption)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{SERVICE_NAME} 사용자 매뉴얼 | ")
    set_run_font(run, size=9, color=CAPTION_COLOR)
    add_page_number(footer)

    doc.core_properties.title = DOC_TITLE
    doc.core_properties.subject = SERVICE_NAME
    doc.core_properties.author = ORG_NAME
    enable_field_update_on_open(doc)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. 문서 정보 및 변경 이력")
    add_table(
        doc,
        ["문서 버전", "변경일", "변경 내용", "작성자", "검토자"],
        [[DOC_VERSION, DOC_DATE.isoformat(), "초안 작성", ORG_NAME, "서비스 담당자"]],
        [2.6, 3.0, 8.2, 3.0, 3.0],
    )

    add_heading(doc, "2. 목차")
    add_paragraph(doc, "아래 목차는 Word에서 문서를 열 때 필드를 업데이트하면 페이지 번호가 반영됩니다.")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    add_heading(doc, "3. 개요")
    add_paragraph(doc, "이 문서는 최초 사용자가 화면에 진입해 주요 기능을 이해하고 결과를 확인할 수 있도록 안내합니다.")
    add_bullets(
        doc,
        [
            "서비스 목적: 사용자가 화면에서 제공하는 기능을 쉽게 이해하고 사용할 수 있도록 지원합니다.",
            "매뉴얼 목적: 접속, 입력, 진행 확인, 결과 활용, 오류 대응까지 따라 할 수 있도록 안내합니다.",
            "기능 범위: 주요 화면 구성, 대표 사용 흐름, 결과 확인, 오류 대응, 자주 묻는 질문을 포함합니다.",
        ],
    )

    add_heading(doc, "4. 대상 사용자")
    add_table(
        doc,
        ["사용자 역할", "주요 사용 목적", "권장 사용 범위"],
        [
            ["일반 사용자", "화면에서 주요 기능을 실행하고 결과를 확인합니다.", "기본 사용 흐름"],
            ["검토자", "결과를 확인하고 필요한 의견을 남깁니다.", "결과 검토 및 피드백"],
        ],
        [3.0, 7.0, 7.0],
    )

    add_heading(doc, "5. 주요 사용 흐름")
    add_scenario(
        doc,
        "5.1 대표 기능 사용하기",
        "가장 일반적인 흐름으로 입력, 진행 확인, 결과 확인까지 이어지는 사용 과정을 안내합니다.",
        ["서비스에 접속할 수 있는 계정과 권한이 있습니다.", "필요한 기본 데이터 또는 설정이 준비되어 있습니다."],
        [
            "브라우저에서 서비스에 접속합니다.",
            "화면의 입력 영역에 사용 목적에 맞는 내용을 입력합니다.",
            "실행 버튼을 눌러 처리를 시작합니다.",
            "진행 표시를 확인하며 처리가 완료될 때까지 기다립니다.",
            "결과 영역에서 생성된 결과와 활용 버튼을 확인합니다.",
        ],
        [
            "진행 상태가 화면에 표시됩니다.",
            "최종 결과 영역이 표시됩니다.",
            "결과를 복사하거나 다음 작업으로 이어갈 수 있습니다.",
        ],
        "result",
        "그림 1. 대표 기능의 최종 결과 화면",
    )

    add_heading(doc, "6. 화면 항목 및 옵션 설명")
    add_table(
        doc,
        ["화면 항목/옵션", "사용 여부", "설명"],
        [
            ["입력 영역", "필수", "사용자가 처리할 내용을 입력합니다."],
            ["실행 버튼", "필수", "입력한 내용을 기준으로 처리를 시작합니다."],
            ["결과 영역", "자동 표시", "완료 후 결과와 후속 동작을 표시합니다."],
        ],
        [4.0, 3.0, 10.0],
    )

    add_heading(doc, "7. FAQ / 자주 묻는 질문")
    add_table(
        doc,
        ["질문", "답변"],
        [
            ["처리가 오래 걸리면 어떻게 하나요?", "진행 상태를 확인하고 필요하면 조건을 줄여 다시 시도합니다."],
            ["결과가 없으면 어떻게 하나요?", "입력 조건을 넓히거나 다른 기준으로 다시 입력합니다."],
        ],
        [6.0, 12.0],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
