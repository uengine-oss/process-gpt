"""Generate the end-user manual DOCX for completion_agent-memory-chat."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SPEC_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = SPEC_ROOT / "docs" / "completion_agent-memory-chat-user-manual.docx"
SCREENSHOT_DIR = SPEC_ROOT / "e2e" / "results" / "screenshots"

DOC_VERSION = "v1.0"
DOC_DATE = date.today()
ORG_NAME = "유엔진솔루션즈"
SERVICE_NAME = "에이전트 메모리 대화"
DOC_TITLE = "에이전트 메모리 대화 사용자 매뉴얼"
FONT_NAME = "Malgun Gothic"
PRIMARY_COLOR = "1F4E79"
CAPTION_COLOR = "666666"
TABLE_HEADER_FILL = "D9EAF7"


SCREENSHOTS = {
    "learning-initial": "process-gpt-completion_agent-memory-chat-01-learning-initial.png",
    "learning-input": "process-gpt-completion_agent-memory-chat-01-learning-input.png",
    "learning-result": "process-gpt-completion_agent-memory-chat-01-learning-result.png",
    "dup-first": "process-gpt-completion_agent-memory-chat-02-learning-first-stored.png",
    "dup-input": "process-gpt-completion_agent-memory-chat-02-learning-duplicate-input.png",
    "dup-skip": "process-gpt-completion_agent-memory-chat-02-learning-duplicate-skip.png",
    "query-initial": "process-gpt-completion_agent-memory-chat-03-query-initial.png",
    "query-input": "process-gpt-completion_agent-memory-chat-03-query-input.png",
    "query-result": "process-gpt-completion_agent-memory-chat-03-query-result.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차는 Word에서 필드 업데이트 시 페이지 번호와 함께 표시됩니다."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, bold=True, color=PRIMARY_COLOR)


def add_paragraph(doc: Document, text: str = "", style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        run = paragraph.add_run(f"{idx}.\t{item}")
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, TABLE_HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_image(doc: Document, key: str, caption: str) -> None:
    path = SCREENSHOT_DIR / SCREENSHOTS[key]
    if not path.exists():
        add_paragraph(doc, f"[화면 예시 누락] {path.name}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=CAPTION_COLOR)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SERVICE_NAME)
    set_run_font(run, size=18, bold=True, color=PRIMARY_COLOR)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DOC_TITLE)
    set_run_font(run, size=28, bold=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 버전", DOC_VERSION],
            ["작성일", DOC_DATE.isoformat()],
            ["작성자/조직", ORG_NAME],
            ["문서 구분", "사용자 참조 매뉴얼"],
            ["대상 서비스", SERVICE_NAME],
        ],
        [4.0, 12.0],
    )
    doc.add_page_break()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{SERVICE_NAME} 사용자 매뉴얼 | ")
    set_run_font(run, size=9, color=CAPTION_COLOR)
    add_page_number(footer)

    doc.core_properties.title = DOC_TITLE
    doc.core_properties.subject = SERVICE_NAME
    doc.core_properties.author = ORG_NAME
    enable_field_update_on_open(doc)


def add_scenario(
    doc: Document,
    title: str,
    purpose: str,
    preconditions: list[str],
    steps: list[str],
    expected: list[str],
    images: list[tuple[str, str]] | None = None,
) -> None:
    add_heading(doc, title, level=2)
    add_paragraph(doc, f"이럴 때 사용합니다: {purpose}")
    add_paragraph(doc, "사용 전 확인")
    add_bullets(doc, preconditions)
    add_paragraph(doc, "따라 하기")
    add_numbers(doc, steps)
    add_paragraph(doc, "화면에서 확인할 내용")
    add_bullets(doc, expected)
    if images:
        for key, caption in images:
            add_image(doc, key, caption)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    # 1. 문서 정보 및 변경 이력
    add_heading(doc, "1. 문서 정보 및 변경 이력")
    add_table(
        doc,
        ["문서 버전", "변경일", "변경 내용", "작성자", "검토자"],
        [[DOC_VERSION, DOC_DATE.isoformat(), "초안 작성", ORG_NAME, "서비스 담당자"]],
        [2.6, 3.0, 8.2, 3.0, 3.0],
    )

    # 2. 목차
    add_heading(doc, "2. 목차")
    add_paragraph(doc, "아래 목차는 Word에서 문서를 열 때 필드를 업데이트하면 페이지 번호가 자동으로 반영됩니다.")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    # 3. 개요
    add_heading(doc, "3. 개요")
    add_paragraph(
        doc,
        "에이전트 메모리 대화는 사용자가 업무 지식이나 정책을 에이전트에게 학습시키고, 이후 같은 에이전트에게 질문하여 학습된 내용을 활용한 답변을 받을 수 있는 대화 화면입니다.",
    )
    add_bullets(
        doc,
        [
            "서비스 목적: 자주 참고하는 사내 정보·정책을 에이전트에 저장하고 필요할 때 자연어 질문으로 빠르게 확인합니다.",
            "매뉴얼 목적: 화면 접속, 학습 정보 입력, 중복 정보 확인, 질의 응답 확인까지 따라 할 수 있도록 안내합니다.",
            "기능 범위: 학습 탭과 질문 탭의 기본 사용 흐름, 화면 항목 설명, 결과 확인, 오류 및 예외 상황 안내를 포함합니다.",
        ],
    )

    # 4. 대상 사용자
    add_heading(doc, "4. 대상 사용자")
    add_table(
        doc,
        ["사용자 역할", "주요 사용 목적", "권장 사용 범위"],
        [
            ["일반 사용자", "에이전트에 업무 정보를 학습시키고, 필요한 정보를 질문해 답변을 받습니다.", "학습 탭과 질문 탭 사용 전반"],
            ["지식 관리자", "조직 공통 정보를 정리하여 에이전트 메모리를 일관성 있게 유지합니다.", "학습 정보 정리 및 중복 관리"],
            ["검토자", "에이전트가 반환한 답변이 의도한 정보와 일치하는지 확인합니다.", "질의 결과 검토"],
        ],
        [3.2, 8.0, 6.0],
    )

    # 5. 사용 전 확인 사항
    add_heading(doc, "5. 사용 전 확인 사항")
    add_bullets(
        doc,
        [
            "서비스에 로그인할 수 있는 사용자 계정과 비밀번호가 준비되어 있습니다.",
            "사용할 에이전트가 이미 등록되어 있고, 해당 에이전트 화면에 접근할 권한이 있습니다.",
            "최신 버전의 데스크톱 브라우저(Chrome, Edge 등)에서 접속합니다.",
            "학습으로 저장할 정보가 사실에 기반한 내용인지 미리 확인합니다.",
            "질문할 때 어떤 주제·범위로 답변을 받고 싶은지 미리 정리해 두면 결과 품질이 좋아집니다.",
        ],
    )

    # 6. 시작하기
    add_heading(doc, "6. 시작하기")
    add_numbers(
        doc,
        [
            "브라우저에서 서비스 주소로 접속한 뒤 로그인 화면에서 이메일과 비밀번호를 입력하고 로그인합니다.",
            "메뉴 또는 에이전트 목록에서 사용할 메모리 에이전트를 선택하여 에이전트 대화 화면을 엽니다.",
            "화면 좌측의 탭에서 정보를 가르치려면 ‘학습’ 탭, 질문하려면 ‘질문’ 탭을 선택합니다.",
            "하단 채팅 입력창에 내용을 입력하고 Enter 키를 눌러 전송합니다.",
            "전송 후 잠시 기다리면 사용자 메시지와 에이전트 응답 메시지가 화면에 순서대로 표시됩니다.",
        ],
    )

    # 7. 화면 구성
    add_heading(doc, "7. 화면 구성")
    add_image(doc, "learning-initial", "그림 1. 에이전트 메모리 대화 학습 모드 초기 화면")
    add_table(
        doc,
        ["화면 영역", "설명"],
        [
            ["탭 영역(학습/질문)", "정보를 가르치는 학습 모드와 질문하는 질의 모드를 전환합니다."],
            ["대화 영역", "사용자가 보낸 메시지와 에이전트의 응답 메시지가 시간 순으로 표시됩니다."],
            ["채팅 입력창", "학습할 내용 또는 질문을 입력하는 영역입니다."],
            ["전송", "Enter 키 또는 전송 버튼으로 입력한 메시지를 에이전트에 전달합니다."],
            ["에이전트 응답 메시지", "저장 완료, 중복 안내, 검색 기반 답변 등 상황별 안내가 표시됩니다."],
        ],
        [4.5, 12.5],
    )

    # 8. 주요 사용 흐름
    add_heading(doc, "8. 주요 사용 흐름")

    add_scenario(
        doc,
        "8.1 에이전트에 새 정보 학습시키기",
        "업무 정책, 자주 묻는 내용, 사내 규정 등 사용자가 알려주고 싶은 정보를 에이전트의 메모리에 저장합니다.",
        [
            "에이전트 대화 화면에 접속해 있습니다.",
            "저장할 정보가 한두 문장으로 명확하게 정리되어 있습니다.",
        ],
        [
            "좌측 탭에서 ‘학습’ 탭을 선택해 학습 모드로 전환합니다.",
            "채팅 입력창에 저장할 내용을 자연스러운 문장으로 입력합니다. 예: ‘우리 회사 휴가 정책은 연 15일이다’.",
            "Enter 키를 눌러 메시지를 전송합니다.",
            "잠시 후 에이전트가 ‘기억했다’는 의미의 응답을 보내는지 확인합니다.",
        ],
        [
            "대화 영역에 사용자가 입력한 학습 메시지가 먼저 표시됩니다.",
            "이어서 에이전트의 저장 완료 응답 메시지가 표시됩니다.",
            "응답에는 ‘기억’, ‘저장’, ‘학습’과 같은 단어가 포함되어 정보가 메모리에 반영되었음을 알려줍니다.",
        ],
        [
            ("learning-input", "그림 2. 학습할 정보를 입력한 직후의 화면"),
            ("learning-result", "그림 3. 학습 정보가 저장되었음을 알리는 응답 화면"),
        ],
    )

    add_scenario(
        doc,
        "8.2 비슷한 정보를 다시 보내면 중복으로 안내받기",
        "이미 저장된 정보와 같은 주제의 학습 메시지를 다시 보냈을 때, 에이전트가 중복 정보로 인식하여 새로 저장하지 않는 상황을 확인합니다.",
        [
            "8.1 흐름으로 같은 주제(예: 휴가 정책)의 정보를 한 번 이상 학습시켜 둔 상태입니다.",
            "현재 학습 탭이 선택되어 있습니다.",
        ],
        [
            "이전에 학습한 내용과 같은 주제의 새로운 문장을 작성합니다. 예: ‘우리 회사 연차는 15일로 운영한다’.",
            "채팅 입력창에 문장을 입력합니다.",
            "Enter 키로 메시지를 전송합니다.",
            "에이전트 응답이 ‘이미 비슷한 내용이 저장되어 있다’는 의미의 안내인지 확인합니다.",
        ],
        [
            "대화 영역에 두 번째 학습 메시지가 사용자 버블로 표시됩니다.",
            "에이전트 응답은 새로 저장했다는 안내가 아니라, ‘비슷한 내용이 이미 있어 새로 저장하지 않았다’는 의미의 안내로 표시됩니다.",
            "첫 번째 메시지의 저장 완료 응답과 두 번째 메시지의 중복 안내 응답이 시간 순으로 구분되어 보입니다.",
        ],
        [
            ("dup-first", "그림 4. 첫 학습 정보가 정상적으로 저장된 화면"),
            ("dup-input", "그림 5. 유사한 학습 정보를 다시 입력한 직후의 화면"),
            ("dup-skip", "그림 6. 중복으로 안내되어 새로 저장되지 않은 응답 화면"),
        ],
    )

    add_scenario(
        doc,
        "8.3 학습한 내용으로 질문하고 답변 받기",
        "학습 모드로 저장해 둔 정보를 질문 모드에서 자연어로 물어보고, 메모리를 활용한 답변을 받습니다.",
        [
            "8.1 흐름으로 질문 대상 정보가 메모리에 저장되어 있습니다.",
            "학습한 내용과 관련된 질문이 한 문장으로 준비되어 있습니다.",
        ],
        [
            "좌측 탭에서 ‘질문’ 탭을 클릭해 질의 모드로 전환합니다.",
            "채팅 입력창에 질문을 입력합니다. 예: ‘우리 회사 휴가는 며칠인가요?’.",
            "Enter 키로 질문을 전송합니다.",
            "잠시 후 에이전트가 검색 기반의 답변 메시지를 보내는지 확인합니다.",
        ],
        [
            "대화 영역에 사용자의 질문이 표시되고, 이어서 에이전트의 답변 메시지가 새 버블로 추가됩니다.",
            "답변에는 학습해 둔 정보(예: 휴가 일수)가 자연스러운 문장으로 포함됩니다.",
            "질문 탭에서는 학습 응답이 아닌 검색 기반 답변 형식으로 결과가 표시됩니다.",
        ],
        [
            ("query-initial", "그림 7. 질의 모드 초기 화면"),
            ("query-input", "그림 8. 질문 메시지를 입력한 직후의 화면"),
            ("query-result", "그림 9. 메모리를 활용한 답변이 표시된 화면"),
        ],
    )

    # 9. 화면 항목 및 옵션 설명
    add_heading(doc, "9. 화면 항목 및 옵션 설명")
    add_table(
        doc,
        ["화면 항목/옵션", "사용 여부", "기본값", "설명 및 예시"],
        [
            ["학습/질문 탭", "필수 선택", "이전 선택 유지", "정보를 저장하려면 학습, 정보를 묻고자 하면 질문 탭을 선택합니다."],
            ["채팅 입력창", "필수", "비어 있음", "한 문장 이상 자연어로 입력합니다. 예: ‘우리 회사 휴가 정책은 연 15일이다’."],
            ["전송(Enter)", "필수", "-", "입력한 내용을 에이전트에게 전송합니다."],
            ["대화 영역", "자동 표시", "-", "사용자 메시지와 에이전트 응답이 시간 순으로 누적되어 표시됩니다."],
        ],
        [4.5, 3.0, 3.0, 7.5],
    )

    # 10. 결과 확인
    add_heading(doc, "10. 결과 확인")
    add_bullets(
        doc,
        [
            "학습 모드에서는 에이전트가 ‘기억/저장/학습’ 의미의 응답을 보내면 정보가 정상적으로 저장된 것입니다.",
            "같은 주제의 정보를 다시 보내 ‘이미 비슷한 내용이 있다’는 안내가 보이면, 새로 저장하지 않은 것이며 기존 정보가 그대로 유지됩니다.",
            "질문 모드에서는 답변 메시지에 학습해 둔 정보가 자연어로 포함되어야 합니다.",
            "원하는 내용과 다른 답변이 나오는 경우, 학습한 표현과 가까운 단어로 질문을 다시 작성해 보면 결과가 개선됩니다.",
            "대화 내용은 화면을 새로 고침해도 동일 에이전트와 사용자 조합에서는 이어서 확인할 수 있습니다.",
        ],
    )

    # 11. 오류 및 예외 상황
    add_heading(doc, "11. 오류 및 예외 상황")
    add_table(
        doc,
        ["상황", "원인", "해결 방법"],
        [
            [
                "메시지를 보내도 응답이 오지 않습니다.",
                "네트워크 일시 장애 또는 에이전트 서비스 일시 중단",
                "잠시 후 같은 메시지를 다시 보내거나 화면을 새로 고침한 뒤 재시도합니다.",
            ],
            [
                "학습을 보냈는데 ‘이미 비슷한 내용이 있다’는 안내가 나옵니다.",
                "같은 주제의 정보가 이전에 이미 저장되어 있습니다.",
                "정상 동작입니다. 새로 저장하려면 기존 정보와 분명히 다른 내용으로 다시 작성합니다.",
            ],
            [
                "질문을 했는데 학습한 내용이 답변에 반영되지 않습니다.",
                "질문이 학습한 표현과 멀거나, 학습 정보가 아직 저장되지 않았을 수 있습니다.",
                "학습 탭에서 정보를 다시 한 번 저장하거나, 학습 시 사용한 단어와 가까운 표현으로 질문을 다시 작성합니다.",
            ],
            [
                "잘못된 에이전트 화면에 들어와 응답이 의도와 다릅니다.",
                "다른 에이전트의 메모리는 공유되지 않습니다.",
                "메뉴에서 사용할 에이전트를 다시 선택해 정확한 에이전트 대화 화면으로 이동합니다.",
            ],
            [
                "에이전트 화면에 접근할 수 없습니다.",
                "해당 에이전트에 대한 접근 권한이 없습니다.",
                "조직 관리자에게 에이전트 사용 권한 부여를 요청합니다.",
            ],
        ],
        [4.5, 6.0, 7.5],
    )

    # 12. 권한 및 역할별 기능 차이
    add_heading(doc, "12. 권한 및 역할별 기능 차이")
    add_table(
        doc,
        ["역할", "접근 가능 화면", "수행 가능한 작업", "제한 사항"],
        [
            ["일반 사용자", "권한이 있는 에이전트의 대화 화면", "학습 정보 저장, 질의 응답 확인", "권한이 없는 다른 에이전트의 메모리에는 접근할 수 없습니다."],
            ["지식 관리자", "조직에서 공유하는 에이전트 대화 화면", "공통 정보의 학습 및 정리, 중복 정보 확인", "에이전트 자체의 환경 설정은 별도 화면에서 권한 필요"],
            ["검토자", "공유된 에이전트의 대화 화면", "질의 응답 결과 확인 및 검토", "학습 정보 신규 입력 권한이 제한될 수 있습니다(권한 필요)."],
        ],
        [3.2, 5.0, 5.5, 5.0],
    )

    # 13. FAQ
    add_heading(doc, "13. FAQ / 자주 묻는 질문")
    add_table(
        doc,
        ["질문", "답변"],
        [
            [
                "학습한 정보를 다른 에이전트도 사용할 수 있나요?",
                "아니요. 학습한 정보는 메시지를 보낸 에이전트의 메모리에만 저장되며, 다른 에이전트에서는 검색되지 않습니다.",
            ],
            [
                "이미 저장된 내용을 수정하거나 삭제하려면 어떻게 하나요?",
                "학습 모드에서는 새로운 정보 저장만 안내됩니다. 기존 정보의 수정·삭제가 필요하면 조직 관리자에게 문의합니다.",
            ],
            [
                "왜 학습한 정보와 같은 답변이 한 번에 나오지 않을까요?",
                "질문에서 사용한 단어가 학습한 문장과 멀면 검색이 정확하지 않을 수 있습니다. 학습 시 사용한 단어를 포함해 다시 질문해 보세요.",
            ],
            [
                "여러 사람이 같은 에이전트를 사용해도 되나요?",
                "권한이 있는 사용자라면 동일 에이전트의 메모리를 공유할 수 있습니다. 단, 학습 내용은 모든 사용자가 동일하게 활용하므로 사실 기반의 정보만 저장하기를 권장합니다.",
            ],
            [
                "응답을 받지 못했는데 같은 메시지를 다시 보내도 되나요?",
                "네. 같은 메시지를 다시 보내도 중복 정보로 안내되어 같은 내용이 두 번 저장되지는 않습니다.",
            ],
        ],
        [6.0, 12.0],
    )

    # 14. 효과적인 질문 작성 팁
    add_heading(doc, "14. 효과적인 질문 작성 팁")
    add_bullets(
        doc,
        [
            "학습한 표현을 그대로 포함해 질문하면 더 정확한 답변을 받을 수 있습니다. 예: 학습 시 ‘휴가 정책’이라고 썼다면, 질문도 ‘휴가 정책은 어떻게 되나요?’로 작성합니다.",
            "한 번에 한 가지만 묻습니다. 여러 주제를 한 문장에 섞으면 답변이 모호해질 수 있습니다.",
            "원하는 답변 형식이 있다면 질문에 함께 적어둡니다. 예: ‘숫자만 알려줘’, ‘짧게 요약해줘’.",
            "결과가 만족스럽지 않으면 같은 의미를 다른 단어로 바꾸어 다시 질문합니다.",
            "사실 정보가 아니라 의견·예측을 묻는 질문은 메모리 기반 답변에 적합하지 않을 수 있습니다.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
