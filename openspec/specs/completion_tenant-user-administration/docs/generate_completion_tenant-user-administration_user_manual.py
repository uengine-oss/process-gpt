"""DOCX user manual generator for completion_tenant-user-administration.

Built from the STYLE_REFERENCE bundled with the docx-user-manual skill.
Reuses cover, heading, table, screenshot, caption, TOC, and footer helpers.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SPEC_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = SPEC_ROOT / "docs" / "completion_tenant-user-administration-user-manual.docx"
SCREENSHOT_DIR = SPEC_ROOT / "e2e" / "results" / "screenshots"

DOC_VERSION = "v1.0"
DOC_DATE = date(2026, 5, 27)
ORG_NAME = "uEngine"
SERVICE_NAME = "Process-GPT"
DOC_TITLE = "테넌트·사용자 관리 사용자 매뉴얼"
FONT_NAME = "Malgun Gothic"
PRIMARY_COLOR = "1F4E79"
CAPTION_COLOR = "666666"
TABLE_HEADER_FILL = "D9EAF7"


SCREENSHOTS = {
    "account-settings-initial": "process-gpt-completion_tenant-user-administration-02-account-settings-initial.png",
    "invite-input": "process-gpt-completion_tenant-user-administration-02-invite-input.png",
    "invite-response": "process-gpt-completion_tenant-user-administration-02-invite-response.png",
    "initial-setting-initial": "process-gpt-completion_tenant-user-administration-03-initial-setting-initial.png",
    "initial-setting-input": "process-gpt-completion_tenant-user-administration-03-initial-setting-input.png",
    "initial-setting-response": "process-gpt-completion_tenant-user-administration-03-initial-setting-response.png",
    "account-tab-initial": "process-gpt-completion_tenant-user-administration-04-account-tab-initial.png",
    "account-tab-input": "process-gpt-completion_tenant-user-administration-04-account-tab-input.png",
    "account-tab-response": "process-gpt-completion_tenant-user-administration-04-account-tab-response.png",
    "login-input": "process-gpt-completion_tenant-user-administration-05-login-input.png",
    "tenant-manage-response": "process-gpt-completion_tenant-user-administration-05-tenant-manage-response.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차는 Word에서 필드 업데이트 시 페이지 번호와 함께 표시됩니다."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, bold=True, color=PRIMARY_COLOR)


def add_paragraph(doc: Document, text: str = "", style=None) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_bullets(doc: Document, items) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items) -> None:
    for idx, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        run = paragraph.add_run(f"{idx}.\t{item}")
        set_run_font(run)


def add_table(doc: Document, headers, rows, widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, TABLE_HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_image(doc: Document, key: str, caption: str) -> None:
    path = SCREENSHOT_DIR / SCREENSHOTS[key]
    if not path.exists():
        add_paragraph(doc, f"[화면 예시 누락] {path.name}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.0))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=CAPTION_COLOR)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SERVICE_NAME)
    set_run_font(run, size=18, bold=True, color=PRIMARY_COLOR)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DOC_TITLE)
    set_run_font(run, size=28, bold=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 버전", DOC_VERSION],
            ["작성일", DOC_DATE.isoformat()],
            ["작성자/조직", ORG_NAME],
            ["문서 구분", "사용자 참조 매뉴얼"],
            ["대상 서비스", SERVICE_NAME],
        ],
        [4.0, 12.0],
    )
    doc.add_page_break()


def add_scenario(
    doc: Document,
    title: str,
    purpose: str,
    preconditions,
    steps,
    expected,
    images=None,
) -> None:
    add_heading(doc, title, level=2)
    add_paragraph(doc, f"이럴 때 사용합니다: {purpose}")
    add_paragraph(doc, "사용 전 확인")
    add_bullets(doc, preconditions)
    add_paragraph(doc, "따라 하기")
    add_numbers(doc, steps)
    add_paragraph(doc, "화면에서 확인할 내용")
    add_bullets(doc, expected)
    if images:
        for key, caption in images:
            add_image(doc, key, caption)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{SERVICE_NAME} 테넌트·사용자 관리 매뉴얼 | ")
    set_run_font(run, size=9, color=CAPTION_COLOR)
    add_page_number(footer)

    doc.core_properties.title = DOC_TITLE
    doc.core_properties.subject = SERVICE_NAME
    doc.core_properties.author = ORG_NAME
    enable_field_update_on_open(doc)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. 문서 정보 및 변경 이력")
    add_table(
        doc,
        ["문서 버전", "변경일", "변경 내용", "작성자", "검토자"],
        [[DOC_VERSION, DOC_DATE.isoformat(), "초안 작성", ORG_NAME, "서비스 담당자"]],
        [2.6, 3.0, 8.2, 3.0, 3.0],
    )

    add_heading(doc, "2. 목차")
    add_paragraph(doc, "아래 목차는 Word에서 문서를 열 때 필드를 업데이트하면 페이지 번호가 반영됩니다.")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    add_heading(doc, "3. 개요")
    add_paragraph(
        doc,
        "이 매뉴얼은 Process-GPT 의 테넌트·사용자 관리 기능을 처음 사용하는 운영자와 사용자가 화면에서 계정 초대, 초기 정보 입력, 본인 정보 갱신, 로그인 후 테넌트 적용을 따라 할 수 있도록 안내합니다.",
    )
    add_bullets(
        doc,
        [
            "서비스 목적: 운영자가 사용자를 초대·관리하고, 사용자가 본인 계정과 테넌트 정보를 화면에서 직접 갱신하도록 지원합니다.",
            "매뉴얼 목적: 화면 접속, 입력, 저장, 결과 확인, 오류 대응까지 한 번에 따라 할 수 있게 안내합니다.",
            "기능 범위: 사용자 초대, 초기 정보 설정, 본인 계정 정보 갱신, 로그인 후 테넌트 적용 및 테넌트 관리 화면 진입을 다룹니다.",
        ],
    )

    add_heading(doc, "4. 대상 사용자")
    add_table(
        doc,
        ["사용자 역할", "주요 사용 목적", "권장 사용 범위"],
        [
            ["운영자(어드민)", "사용자를 초대하고 테넌트 정보를 관리합니다.", "사용자 초대, 테넌트 관리 화면"],
            ["초대 받은 사용자", "초대 메일을 통해 초기 사용자명·비밀번호를 설정합니다.", "초기 설정 화면"],
            ["일반 사용자", "본인 계정 정보를 확인하고 사용자명을 갱신합니다.", "계정 설정 Account 탭"],
        ],
        [4.0, 7.0, 6.0],
    )

    add_heading(doc, "5. 사용 전 확인 사항")
    add_bullets(
        doc,
        [
            "사용 가능한 Process-GPT 계정과 비밀번호가 있어야 합니다. (어드민 작업은 어드민 권한이 부여된 계정이 필요합니다.)",
            "최신 버전의 데스크톱 브라우저(Chrome, Edge 등)에서 접속을 권장합니다.",
            "사용자 초대를 진행할 경우, 초대할 상대방의 이메일 주소를 미리 준비합니다.",
            "초기 설정 화면을 이용하려면 초대 메일 또는 사전 안내를 통해 화면 링크가 제공되어 있어야 합니다.",
            "테넌트 관리 화면은 해당 테넌트에 대해 적절한 권한이 부여된 계정에서만 정상적으로 표시됩니다.",
        ],
    )

    add_heading(doc, "6. 시작하기")
    add_paragraph(doc, "처음 접속할 때는 다음 순서를 따라 진행합니다.")
    add_numbers(
        doc,
        [
            "브라우저 주소창에 Process-GPT 접속 주소를 입력하여 로그인 화면을 엽니다.",
            "이메일과 비밀번호를 입력하고 \"Remember me\" 항목을 확인한 뒤 로그인 버튼을 누릅니다.",
            "로그인이 완료되면 사용 중인 테넌트가 자동으로 적용되어 메인 화면으로 이동합니다.",
            "운영자는 우측 상단 메뉴 또는 사용자 메뉴에서 \"계정 설정\" 화면(/account-settings)으로 이동할 수 있습니다.",
            "테넌트 관리 화면이 필요하면 별도 메뉴 또는 안내된 주소(/tenant/manage)로 이동합니다.",
        ],
    )

    add_heading(doc, "7. 화면 구성")
    add_paragraph(doc, "테넌트·사용자 관리에서 자주 사용하는 화면 영역은 다음과 같습니다.")
    add_table(
        doc,
        ["화면", "주요 영역", "역할"],
        [
            ["로그인 화면", "이메일 입력, 비밀번호 입력, Remember me, 로그인 버튼", "사용자 인증과 테넌트 자동 적용의 시작점"],
            ["계정 설정 화면 – Account 탭", "사용자명·이메일 표시 영역, 저장 버튼", "본인 계정 정보 확인 및 사용자명 갱신"],
            ["계정 설정 화면 – ManageAccess 탭", "사용자 목록, \"사용자 추가\" 버튼, 초대 다이얼로그", "사용자 초대 및 관리"],
            ["초기 설정 화면", "사용자명 입력, 새 비밀번호 입력, 비밀번호 확인 입력, \"초기 설정 완료\" 버튼", "초대 받은 사용자가 첫 비밀번호와 사용자명을 등록"],
            ["테넌트 관리 화면", "소유 테넌트 목록, 테넌트 정보 카드", "로그인 후 적용된 테넌트 현황 확인"],
        ],
        [4.5, 6.5, 6.0],
    )
    add_image(doc, "account-settings-initial", "그림 1. 로그인 후 계정 설정 화면 진입 상태")

    add_heading(doc, "8. 주요 사용 흐름")

    add_scenario(
        doc,
        "8.1 사용자 관리 화면에서 신규 사용자 초대하기",
        "운영자가 새로운 동료 또는 사용자에게 Process-GPT 접속 안내 메일을 발송하려고 할 때 사용합니다.",
        [
            "어드민 권한을 가진 계정으로 로그인되어 있습니다.",
            "초대할 사용자의 이메일 주소를 준비합니다.",
            "동일한 이메일이 이미 등록되어 있지 않은지 확인합니다.",
        ],
        [
            "로그인 후 \"계정 설정\" 화면으로 이동합니다.",
            "상단 탭에서 \"ManageAccess(사용자 관리)\" 탭을 선택합니다.",
            "사용자 목록 위의 \"사용자 추가\" 버튼을 누릅니다.",
            "초대 다이얼로그가 열리면 초대할 이메일 주소를 입력합니다.",
            "\"초대\" 버튼을 눌러 안내 메일 발송과 사용자 등록 처리를 시작합니다.",
        ],
        [
            "초대 다이얼로그가 성공 상태로 바뀌거나 사용자 목록에 새 사용자가 표시됩니다.",
            "초대 받은 사용자에게 초기 설정 화면 링크가 포함된 안내 이메일이 발송됩니다.",
            "이메일 형식이 잘못되었거나 이미 가입된 사용자인 경우 화면에 안내 메시지가 표시됩니다.",
        ],
        images=[
            ("invite-input", "그림 2. 사용자 초대 정보 입력 화면"),
            ("invite-response", "그림 3. 사용자 초대 완료 결과 화면"),
        ],
    )

    add_scenario(
        doc,
        "8.2 초대 받은 사용자가 초기 정보 설정하기",
        "초대 메일로 받은 링크를 통해 처음 접속한 사용자가 사용자명과 비밀번호를 등록하려고 할 때 사용합니다.",
        [
            "초대 안내 메일에 포함된 \"초기 설정\" 링크를 이용해 화면에 진입할 수 있습니다.",
            "사용할 사용자명을 미리 정해 둡니다.",
            "8자 이상의 새 비밀번호를 준비합니다.",
        ],
        [
            "초대 메일에서 안내된 링크를 눌러 초기 설정 화면을 엽니다.",
            "사용자명 입력칸에 사용할 이름을 입력합니다.",
            "새 비밀번호 입력칸에 비밀번호를 입력합니다.",
            "비밀번호 확인 입력칸에 동일한 비밀번호를 다시 입력합니다.",
            "\"초기 설정 완료\" 버튼을 누릅니다.",
        ],
        [
            "정보가 정상적으로 저장되면 \"설정이 완료되었습니다\" 형식의 안내가 표시되고 로그인 화면으로 이동합니다.",
            "비밀번호가 서로 다르거나 8자 미만이면 화면에 안내 메시지가 표시되어 다시 입력하도록 안내됩니다.",
            "저장 후에는 새 비밀번호와 사용자명으로 로그인할 수 있습니다.",
        ],
        images=[
            ("initial-setting-initial", "그림 4. 초기 설정 화면 초기 상태"),
            ("initial-setting-input", "그림 5. 초기 설정 정보 입력 화면"),
            ("initial-setting-response", "그림 6. 초기 설정 완료 결과 화면"),
        ],
    )

    add_scenario(
        doc,
        "8.3 본인 계정 정보 갱신하기",
        "로그인한 사용자가 본인의 사용자명 등 표시 정보를 직접 수정하려고 할 때 사용합니다.",
        [
            "본인 계정으로 로그인되어 있습니다.",
            "변경할 새 사용자명을 준비합니다.",
        ],
        [
            "\"계정 설정\" 화면으로 이동합니다. 기본적으로 Account 탭이 활성화되어 있습니다.",
            "현재 사용자 정보(이름, 이메일)가 화면에 표시될 때까지 잠시 기다립니다.",
            "사용자명 입력칸의 값을 새 이름으로 변경합니다.",
            "화면 하단의 \"저장\" 버튼을 눌러 변경 사항을 적용합니다.",
        ],
        [
            "저장이 완료되면 화면 하단에 성공 안내가 잠깐 표시됩니다.",
            "잠시 후 화면이 다시 불러와지며 변경된 사용자명이 표시됩니다.",
            "필수 항목이 비어 있는 경우 저장이 진행되지 않고 입력칸 옆에 안내가 표시됩니다.",
        ],
        images=[
            ("account-tab-initial", "그림 7. 계정 설정 Account 탭 초기 상태"),
            ("account-tab-input", "그림 8. 사용자 정보 갱신 입력 화면"),
            ("account-tab-response", "그림 9. 사용자 정보 갱신 결과 화면"),
        ],
    )

    add_scenario(
        doc,
        "8.4 로그인 후 테넌트 적용 및 테넌트 관리 화면 열기",
        "로그인 시 본인이 속한 테넌트가 자동으로 적용되는지 확인하고 테넌트 관리 화면을 열어보려 할 때 사용합니다.",
        [
            "사용 가능한 계정과 비밀번호가 있습니다.",
            "테넌트 관리 화면에 접근할 수 있는 권한이 부여된 계정입니다.",
        ],
        [
            "로그인 화면(/auth/login)에 접속합니다.",
            "이메일과 비밀번호를 입력합니다.",
            "\"Remember me\" 항목을 체크합니다.",
            "로그인 버튼을 눌러 로그인을 진행합니다.",
            "로그인이 완료되어 메인 화면으로 이동한 뒤, 메뉴 또는 안내된 주소(/tenant/manage)를 통해 테넌트 관리 화면으로 이동합니다.",
        ],
        [
            "로그인 직후 사용 중인 테넌트가 자동으로 적용되어 메인 화면이 열립니다.",
            "테넌트 관리 화면이 정상적으로 로드되고, 소유 또는 참여 중인 테넌트 정보가 표시됩니다.",
            "권한이 부족한 계정으로 접근하면 테넌트 관리 화면 내용이 비어 있거나 안내 메시지가 표시될 수 있습니다.",
        ],
        images=[
            ("login-input", "그림 10. 로그인 화면 입력 완료 상태"),
            ("tenant-manage-response", "그림 11. 테넌트 적용 후 테넌트 관리 화면"),
        ],
    )

    add_heading(doc, "9. 화면 항목 및 옵션 설명")
    add_table(
        doc,
        ["화면", "항목", "필수 여부", "설명 및 예시"],
        [
            ["로그인", "이메일", "필수", "가입된 이메일 주소를 입력합니다."],
            ["로그인", "비밀번호", "필수", "본인 계정의 비밀번호를 입력합니다."],
            ["로그인", "Remember me", "필수 확인", "로그인 진행을 위해 체크합니다."],
            ["사용자 초대", "이메일", "필수", "초대할 사용자의 유효한 이메일을 입력합니다. 예: invitee@example.com"],
            ["초기 설정", "사용자명", "필수", "다른 사용자에게 표시되는 이름입니다."],
            ["초기 설정", "새 비밀번호", "필수", "8자 이상으로 설정합니다."],
            ["초기 설정", "비밀번호 확인", "필수", "새 비밀번호와 동일하게 입력합니다."],
            ["Account 탭", "사용자명", "필수", "본인 표시 이름을 입력합니다."],
            ["Account 탭", "이메일", "표시 전용", "현재 계정 이메일을 보여 줍니다."],
        ],
        [3.5, 3.5, 2.5, 8.0],
    )

    add_heading(doc, "10. 결과 확인")
    add_bullets(
        doc,
        [
            "사용자 초대: 초대 다이얼로그가 성공 상태로 바뀌고 사용자 목록 또는 안내 메시지로 결과를 확인할 수 있습니다.",
            "초기 설정 완료: 완료 안내 후 로그인 화면으로 이동하며, 다음부터는 새 비밀번호와 사용자명으로 접속합니다.",
            "본인 정보 갱신: 성공 안내가 잠깐 표시된 뒤 화면이 다시 로드되어 새 사용자명이 반영됩니다.",
            "테넌트 적용: 로그인 후 메인 화면이 열리고, 테넌트 관리 화면에서 사용 중인 테넌트 정보를 확인할 수 있습니다.",
        ],
    )

    add_heading(doc, "11. 오류 및 예외 상황")
    add_table(
        doc,
        ["상황", "원인", "해결 방법"],
        [
            ["로그인 버튼을 눌러도 로그인이 진행되지 않음", "이메일·비밀번호 형식 오류 또는 Remember me 미체크", "입력값을 다시 확인하고 Remember me 항목을 체크한 뒤 다시 시도합니다."],
            ["로그인이 실패하고 안내 메시지가 표시됨", "비밀번호 불일치 또는 계정 비활성화", "비밀번호를 다시 입력하거나 운영자에게 문의합니다."],
            ["사용자 초대가 실패함", "이메일 형식 오류 또는 이미 등록된 사용자", "이메일 주소를 다시 확인하거나 다른 이메일로 시도합니다."],
            ["초기 설정 저장이 진행되지 않음", "비밀번호 8자 미만 또는 비밀번호 확인 불일치", "비밀번호 정책을 충족하도록 다시 입력합니다."],
            ["계정 설정 저장 후 변경 사항이 보이지 않음", "화면이 자동으로 다시 로드되기 전 상태", "잠시 기다리거나 페이지를 새로 고침합니다."],
            ["테넌트 관리 화면이 비어 있음", "해당 계정에 부여된 테넌트가 없거나 권한이 부족함", "운영자에게 테넌트 권한 부여를 요청합니다."],
        ],
        [5.0, 5.0, 7.0],
    )

    add_heading(doc, "12. 권한 및 역할별 기능 차이")
    add_table(
        doc,
        ["기능", "운영자(어드민)", "초대 받은 사용자", "일반 사용자"],
        [
            ["사용자 초대", "사용 가능", "권한 필요", "권한 필요"],
            ["초기 정보 설정", "사용 가능", "사용 가능 (본인 최초 1회)", "권한 필요 (해당 없음)"],
            ["본인 계정 정보 갱신", "사용 가능", "로그인 이후 사용 가능", "사용 가능"],
            ["테넌트 관리 화면 열람", "사용 가능", "권한 필요", "권한 필요"],
        ],
        [4.5, 4.0, 4.0, 4.0],
    )

    add_heading(doc, "13. FAQ / 자주 묻는 질문")
    add_table(
        doc,
        ["질문", "답변"],
        [
            ["초대 메일이 도착하지 않습니다.", "스팸함을 먼저 확인하고, 일정 시간 이상 도착하지 않으면 운영자에게 다시 초대를 요청합니다."],
            ["초기 설정 화면 링크를 잃어버렸습니다.", "운영자에게 다시 초대를 요청하면 새 안내 메일이 발송됩니다."],
            ["비밀번호를 잊었습니다.", "운영자에게 초기화 또는 재초대를 요청합니다."],
            ["사용자명을 바꿔도 화면에 그대로 보입니다.", "저장 후 화면이 다시 로드될 때까지 잠시 기다리거나 페이지를 새로 고침합니다."],
            ["테넌트 관리 화면이 비어 있습니다.", "해당 계정에 부여된 테넌트 권한을 운영자에게 확인합니다."],
            ["같은 이메일로 두 번 초대해도 되나요?", "중복 초대 시 안내 메시지가 표시되므로 다른 이메일로 다시 시도합니다."],
        ],
        [6.0, 11.0],
    )

    add_heading(doc, "14. 효과적인 사용 팁")
    add_bullets(
        doc,
        [
            "사용자 초대 전, 상대방의 정확한 업무 이메일을 미리 확인하면 중복 초대를 줄일 수 있습니다.",
            "초기 설정 시 사용자명을 팀 내부 표기 규칙에 맞춰 통일하면 사용자 관리가 쉬워집니다.",
            "비밀번호는 8자 이상이면서 영문·숫자·기호를 함께 사용하면 보안에 유리합니다.",
            "본인 정보를 자주 갱신할 때는 변경 사항을 한 번에 입력하고 저장하면 화면 새로고침 횟수를 줄일 수 있습니다.",
            "테넌트 관리 화면에서 정보가 보이지 않을 때는 권한 문제일 가능성이 높으므로 운영자에게 문의합니다.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
