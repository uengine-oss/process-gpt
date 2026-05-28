"""워크아이템 자동 진행 사용자 매뉴얼 생성 스크립트.

이 스크립트는 openspec/specs/completion_workitem-polling-execution 의 명세 및
관련 사용자 동선 산출물을 바탕으로 최종 사용자용 DOCX 매뉴얼을 생성합니다.
스타일/헬퍼 패턴은 .claude/skills/docx-user-manual/STYLE_REFERENCE.py 를 기반으로 합니다.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SPEC_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = SPEC_ROOT / "docs"
SCREENSHOT_DIR = SPEC_ROOT / "e2e" / "results" / "screenshots"
OUTPUT_PATH = DOCS_DIR / "completion_workitem-polling-execution-user-manual.docx"

DOC_VERSION = "v1.0"
DOC_DATE = date.today()
ORG_NAME = "process-GPT 운영팀"
SERVICE_NAME = "process-GPT"
DOC_TITLE = "워크아이템 자동 진행 사용자 매뉴얼"
FONT_NAME = "Malgun Gothic"
PRIMARY_COLOR = "1F4E79"
CAPTION_COLOR = "666666"
TABLE_HEADER_FILL = "D9EAF7"


SCREENSHOTS = {
    "login": "process-gpt-completion_workitem-polling-execution-01-login.png",
    "submitted_initial": "process-gpt-completion_workitem-polling-execution-02-submitted-initial.png",
    "claimed_or_done": "process-gpt-completion_workitem-polling-execution-03-claimed-or-done.png",
    "card_detail": "process-gpt-completion_workitem-polling-execution-04-card-detail.png",
}


# ---------------------------------------------------------------------------
# 헬퍼
# ---------------------------------------------------------------------------


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False, align_center: bool = True) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차는 Word에서 필드 업데이트 시 페이지 번호와 함께 표시됩니다."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, bold=True, color=PRIMARY_COLOR)


def add_paragraph(doc: Document, text: str = "", bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items: list[str]) -> None:
    """매 호출마다 1부터 시작하는 번호 매김(스타일 상속 회피)."""
    for idx, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        run = paragraph.add_run(f"{idx}.\t{item}")
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, TABLE_HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, align_center=False)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_image(doc: Document, key: str, caption: str) -> None:
    path = SCREENSHOT_DIR / SCREENSHOTS[key]
    if not path.exists():
        add_paragraph(doc, f"[화면 예시 누락] {path.name}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.0))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=CAPTION_COLOR)


# ---------------------------------------------------------------------------
# 문서 섹션
# ---------------------------------------------------------------------------


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SERVICE_NAME)
    set_run_font(run, size=18, bold=True, color=PRIMARY_COLOR)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DOC_TITLE)
    set_run_font(run, size=28, bold=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 버전", DOC_VERSION],
            ["작성일", DOC_DATE.isoformat()],
            ["작성자/조직", ORG_NAME],
            ["문서 구분", "사용자 참조 매뉴얼"],
            ["대상 서비스", SERVICE_NAME],
        ],
        [4.0, 12.0],
    )
    doc.add_page_break()


def add_scenario(
    doc: Document,
    title: str,
    purpose: str,
    preconditions: list[str],
    steps: list[str],
    expected: list[str],
    screenshot_key: str | None = None,
    caption: str | None = None,
) -> None:
    add_heading(doc, title, level=2)
    add_paragraph(doc, f"이럴 때 사용합니다: {purpose}")
    add_paragraph(doc, "사용 전 확인", bold=True)
    add_bullets(doc, preconditions)
    add_paragraph(doc, "따라 하기", bold=True)
    add_numbers(doc, steps)
    add_paragraph(doc, "화면에서 확인할 내용", bold=True)
    add_bullets(doc, expected)
    if screenshot_key and caption:
        add_image(doc, screenshot_key, caption)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{SERVICE_NAME} {DOC_TITLE} | ")
    set_run_font(run, size=9, color=CAPTION_COLOR)
    add_page_number(footer)

    doc.core_properties.title = DOC_TITLE
    doc.core_properties.subject = SERVICE_NAME
    doc.core_properties.author = ORG_NAME
    enable_field_update_on_open(doc)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    # 1. 문서 정보 및 변경 이력
    add_heading(doc, "1. 문서 정보 및 변경 이력")
    add_table(
        doc,
        ["문서 버전", "변경일", "변경 내용", "작성자", "검토자"],
        [[DOC_VERSION, DOC_DATE.isoformat(), "초안 작성", ORG_NAME, "서비스 담당자"]],
        [2.6, 3.0, 6.4, 3.0, 3.0],
    )

    # 2. 목차
    add_heading(doc, "2. 목차")
    add_paragraph(doc, "아래 목차는 Word에서 문서를 열 때 필드를 업데이트하면 페이지 번호가 반영됩니다.")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    # 3. 개요
    add_heading(doc, "3. 개요")
    add_paragraph(
        doc,
        "이 문서는 process-GPT에서 사용자가 제출하거나 대기 중인 할 일 카드가 시스템에 의해 자동으로 진행되고, "
        "다음 단계의 할 일이 만들어지는 과정을 사용자 화면 관점에서 안내합니다.",
    )
    add_bullets(
        doc,
        [
            "서비스 목적: 사용자가 제출한 할 일을 시스템이 주기적으로 감지해 자동으로 진행하고, 다음 활동을 이어 만듭니다.",
            "매뉴얼 목적: 로그인부터 할 일 제출, 자동 진행 확인, 다음 활동 인계, 오류 대응까지 한 흐름으로 따라 할 수 있도록 안내합니다.",
            "기능 범위: 할 일 보드(칸반), 자동 진행, 다음 활동 생성, 프로세스 종료 상태 확인, 오류·재시도 안내, 외부 고객 안내 메일을 다룹니다.",
        ],
    )

    # 4. 대상 사용자
    add_heading(doc, "4. 대상 사용자")
    add_table(
        doc,
        ["사용자 역할", "주요 사용 목적", "권장 사용 범위"],
        [
            ["일반 사용자", "본인에게 배정된 할 일을 제출하고 자동 진행 상태를 확인합니다.", "할 일 보드 사용, 카드 제출, 결과 확인"],
            ["프로세스 담당자", "프로세스 인스턴스 전체의 진행과 종료 상태를 확인합니다.", "인스턴스 상태 모니터링, 정체된 할 일 확인"],
            ["외부 고객", "안내 이메일의 외부 폼 링크를 통해 본인 작업을 입력합니다.", "안내 이메일 수신 및 외부 폼 입력"],
        ],
        [3.2, 7.0, 7.0],
    )

    # 5. 사용 전 확인 사항
    add_heading(doc, "5. 사용 전 확인 사항")
    add_bullets(
        doc,
        [
            "process-GPT 사용 권한을 가진 계정과 비밀번호가 준비되어 있습니다.",
            "최신 버전의 Chrome, Edge 등 표준 웹 브라우저를 사용합니다.",
            "본인에게 배정된 할 일을 알고 있고, 필요한 입력 정보(폼 항목, 첨부 데이터 등)를 미리 준비합니다.",
            "외부 고객에게 안내가 필요한 경우, 이메일 수신이 가능한 주소가 프로세스 정보에 등록되어 있습니다.",
        ],
    )

    # 6. 시작하기
    add_heading(doc, "6. 시작하기")
    add_paragraph(doc, "처음 사용자는 아래 순서로 화면에 진입합니다.")
    add_numbers(
        doc,
        [
            "웹 브라우저에서 process-GPT 접속 주소로 이동합니다.",
            "로그인 화면에서 사용자 계정과 비밀번호를 입력해 로그인합니다.",
            "로그인 후 첫 화면에서 본인에게 배정된 할 일 보드(칸반 보드)로 이동합니다.",
            "할 일 보드에서 현재 컬럼별 카드 분포를 확인합니다.",
        ],
    )
    add_image(doc, "login", "그림 1. 로그인 직후 첫 진입 화면")

    # 7. 화면 구성
    add_heading(doc, "7. 화면 구성")
    add_paragraph(doc, "할 일 보드와 카드 상세 화면은 다음과 같이 구성되어 있습니다.")
    add_table(
        doc,
        ["화면 영역", "주요 기능", "표시되는 내용"],
        [
            ["할 일 보드(칸반)", "본인에게 배정된 카드 전체 흐름을 한눈에 확인", "제출 대기, 진행 중, 완료 등 상태별 컬럼과 카드"],
            ["제출 대기 컬럼", "사용자가 입력하거나 시스템 처리가 시작되기 전 카드 영역", "최근 등록된 할 일 카드 목록"],
            ["진행 중 컬럼", "시스템 또는 담당자가 처리 중인 카드 영역", "현재 자동 처리 중이거나 입력 대기 중인 카드"],
            ["완료 컬럼", "처리가 끝난 카드 영역", "완료된 활동 카드와 결과"],
            ["카드 상세 영역", "카드 클릭 시 열리는 입력·로그·결과 패널", "활동 이름, 입력 폼, 처리 로그, 결과 메시지"],
            ["상태 배지", "각 카드와 인스턴스의 현재 상태를 표시", "제출됨, 진행 중, 대기, 완료, 오류 등 상태 표시"],
        ],
        [4.5, 5.5, 7.0],
    )
    add_image(doc, "submitted_initial", "그림 2. 제출 대기 컬럼에 새 할 일 카드가 노출된 초기 상태")

    # 8. 주요 사용 흐름
    add_heading(doc, "8. 주요 사용 흐름")

    add_scenario(
        doc,
        "8.1 제출한 할 일이 자동으로 처리되는 것 확인하기",
        "본인이 제출했거나 시스템에 새로 등록된 할 일 카드가 시스템에 의해 자동으로 다음 단계로 넘어가는지 확인하는 가장 기본적인 흐름입니다.",
        [
            "본인 계정으로 로그인되어 있어야 합니다.",
            "할 일 보드의 제출 대기 컬럼에 본인에게 배정된 카드가 한 개 이상 있어야 합니다.",
        ],
        [
            "할 일 보드 화면을 엽니다.",
            "제출 대기 컬럼에 본인 카드가 표시되는지 확인합니다.",
            "잠시 후 페이지를 새로고침하거나 자동 갱신을 기다립니다.",
            "카드가 진행 중 또는 완료 컬럼으로 이동하는 것을 확인합니다.",
            "필요하면 이동한 카드를 클릭해 상세 화면에서 처리 로그를 확인합니다.",
        ],
        [
            "제출 대기 컬럼에서 카드가 사라지고 진행 중 또는 완료 컬럼에 같은 카드가 표시됩니다.",
            "카드 상세에는 시스템이 자동으로 처리했음을 알 수 있는 처리 로그가 남습니다.",
            "다음 단계의 활동이 있다면 새 카드가 진행 중 컬럼에 자동으로 등장합니다.",
        ],
        "claimed_or_done",
        "그림 3. 자동 처리 후 카드가 진행 중·완료 컬럼으로 이동한 상태",
    )

    add_scenario(
        doc,
        "8.2 완료 후 다음 단계 할 일 카드 받기",
        "현재 활동이 정상적으로 끝나면, 다음 단계의 할 일 카드가 자동으로 만들어져 담당자에게 배정됩니다.",
        [
            "프로세스 정의에 다음 활동이 있어야 합니다.",
            "현재 카드의 결과가 다음 단계를 진행할 수 있는 정상 상태여야 합니다.",
        ],
        [
            "완료 컬럼으로 이동한 카드를 클릭해 결과를 확인합니다.",
            "할 일 보드로 돌아와 진행 중 컬럼에 새 카드가 생성되었는지 확인합니다.",
            "새 카드의 담당자가 본인이라면 클릭해 상세를 열어 다음 작업을 이어갑니다.",
            "담당자가 다른 사람이라면 별도 조치 없이 상태를 모니터링합니다.",
        ],
        [
            "이전 카드의 상태 배지가 \"완료\"로 표시됩니다.",
            "다음 활동에 해당하는 새 카드가 진행 중 컬럼에 자동으로 만들어집니다.",
            "프로세스 인스턴스 상태는 아직 끝나지 않았다면 \"진행 중\"으로 유지됩니다.",
        ],
        "card_detail",
        "그림 4. 카드 상세 화면 - 처리 로그와 결과 확인",
    )

    add_scenario(
        doc,
        "8.3 마지막 단계가 끝나 프로세스가 종료되는 것 확인하기",
        "프로세스의 마지막 활동이 완료되면 인스턴스 전체가 종료 상태가 되고, 종료 시각이 기록됩니다.",
        [
            "프로세스 정의에 종료 활동(마지막 단계)이 명시되어 있어야 합니다.",
            "마지막 활동 카드가 완료 컬럼으로 이동한 상태여야 합니다.",
        ],
        [
            "마지막 활동 카드가 완료 컬럼으로 이동했는지 확인합니다.",
            "프로세스 인스턴스 상세 또는 인스턴스 목록 화면으로 이동합니다.",
            "해당 인스턴스의 상태 배지와 종료 시각을 확인합니다.",
        ],
        [
            "인스턴스 상태 배지가 \"완료\"로 표시됩니다.",
            "인스턴스 상세에 종료 시각이 기록되어 있습니다.",
            "할 일 보드의 진행 중 컬럼에서 해당 인스턴스의 카드는 더 이상 등장하지 않습니다.",
        ],
    )

    add_scenario(
        doc,
        "8.4 하위 프로세스 결과를 기다렸다가 본 프로세스 다시 진행하기",
        "현재 활동이 하위 프로세스를 호출하는 단계라면, 시스템이 하위 프로세스를 시작하고 그 결과가 끝날 때까지 본 활동을 대기 상태로 유지합니다. 하위 작업이 모두 끝나면 본 프로세스의 다음 단계 카드가 자동으로 다시 만들어집니다.",
        [
            "현재 활동이 하위 프로세스 호출 단계로 정의되어 있어야 합니다.",
            "하위 프로세스의 담당자가 정해져 있고, 입력 데이터가 전달 가능해야 합니다.",
        ],
        [
            "본인의 카드가 \"대기\" 상태로 표시되는 것을 확인합니다.",
            "하위 프로세스 담당자가 하위 작업을 완료할 때까지 기다립니다.",
            "할 일 보드를 갱신해 본 프로세스의 다음 단계 카드가 다시 등장하는지 확인합니다.",
            "새 카드가 본인 담당이라면 클릭해 작업을 이어갑니다.",
        ],
        [
            "하위 프로세스 진행 중에는 본 카드의 상태가 \"대기\"로 표시됩니다.",
            "하위 작업이 모두 끝나면 본 프로세스의 다음 활동 카드가 새로 등장합니다.",
            "카드 상세의 로그에 하위 프로세스 완료 후 본 작업이 다시 진행됨을 알 수 있는 메시지가 표시됩니다.",
        ],
    )

    add_scenario(
        doc,
        "8.5 외부 고객에게 안내 이메일 발송되는 흐름 확인하기",
        "다음 단계의 담당자가 외부 고객일 때, 시스템이 외부 고객의 이메일 주소로 안내 메일과 외부 폼 링크를 자동으로 발송합니다.",
        [
            "다음 활동의 담당 역할이 외부 고객으로 지정되어 있어야 합니다.",
            "프로세스 정보에 외부 고객의 이메일 주소가 등록되어 있어야 합니다.",
            "이메일 수신함이 정상적으로 동작합니다.",
        ],
        [
            "현재 활동 카드의 입력을 완료하고 제출합니다.",
            "할 일 보드를 갱신해 현재 카드가 완료 컬럼으로 이동하는 것을 확인합니다.",
            "외부 고객에게 시스템 발송 안내 이메일이 도착했는지 확인합니다.",
            "이메일 본문의 외부 폼 링크를 통해 외부 고객이 본인 작업을 입력할 수 있도록 안내합니다.",
        ],
        [
            "현재 카드가 완료 컬럼으로 이동합니다.",
            "외부 고객의 이메일 수신함에 안내 메일과 외부 폼 링크가 도착합니다.",
            "외부 고객이 입력을 완료하면 다음 단계의 진행이 자동으로 이어집니다.",
        ],
    )

    # 9. 화면 항목 및 옵션 설명
    add_heading(doc, "9. 화면 항목 및 옵션 설명")
    add_table(
        doc,
        ["화면 항목/옵션", "사용 여부", "기본값", "설명"],
        [
            ["이메일", "필수", "없음", "로그인 시 사용하는 사용자 계정 이메일 주소입니다."],
            ["비밀번호", "필수", "없음", "사용자 계정 비밀번호입니다."],
            ["할 일 카드 제목", "자동 표시", "활동 이름", "현재 활동의 이름이 카드 제목으로 표시됩니다."],
            ["상태 배지", "자동 표시", "현재 상태", "제출됨, 진행 중, 대기, 완료, 오류 등 카드의 현재 상태를 색상과 함께 표시합니다."],
            ["카드 입력 폼", "조건부 필수", "프로세스 정의에 따름", "사용자가 직접 입력해야 하는 항목이 있을 때만 표시됩니다. 필수 항목은 별도로 표시됩니다."],
            ["제출 버튼", "필수", "비활성", "필수 입력이 모두 완료되면 활성화되어 카드를 제출할 수 있습니다."],
            ["처리 로그 영역", "자동 표시", "비어 있음", "시스템 자동 처리 결과, 재시도 안내, 오류 메시지가 시간 순으로 표시됩니다."],
        ],
        [4.0, 2.5, 3.0, 7.0],
    )

    # 10. 결과 확인
    add_heading(doc, "10. 결과 확인")
    add_bullets(
        doc,
        [
            "완료된 카드는 완료 컬럼에 표시되며, 카드 상세에서 처리 결과와 로그를 확인할 수 있습니다.",
            "다음 단계 활동이 있을 경우, 진행 중 컬럼에 새 카드가 자동으로 등장합니다.",
            "프로세스 전체가 끝나면 인스턴스 상태 배지가 \"완료\"로 바뀌고 종료 시각이 기록됩니다.",
            "완료된 카드는 이후에도 카드 상세에서 입력 내용, 결과, 로그를 다시 확인할 수 있습니다.",
        ],
    )

    # 11. 오류 및 예외 상황
    add_heading(doc, "11. 오류 및 예외 상황")
    add_table(
        doc,
        ["상황", "사용자에게 보이는 내용", "원인", "해결 방법"],
        [
            [
                "자동 진행이 멈춘 것처럼 보임",
                "카드가 제출 대기 또는 진행 중 컬럼에 오래 머무름",
                "시스템이 카드를 처리하는 주기를 기다리는 중이거나, 일시적으로 처리 담당이 정체된 상태",
                "잠시 후 새로고침해 카드 상태가 갱신되는지 확인합니다. 일정 시간이 지나면 시스템이 자동으로 정체 상태를 해제하고 다시 처리합니다.",
            ],
            [
                "카드 상태가 \"대기\"로 표시됨",
                "카드 상세에 \"진행 조건이 충족되지 않았습니다\" 또는 유사한 안내",
                "현재 단계 완료 조건이 충족되지 않았거나, 시스템이 결과를 해석하지 못한 경우",
                "카드 상세의 사유 메시지를 확인하고, 필요한 입력을 보완하거나 담당자에게 문의합니다.",
            ],
            [
                "재시도 안내가 표시됨",
                "로그에 \"다시 시도하겠습니다. (시도 1/3)\" 등 메시지",
                "일시적인 처리 실패로 시스템이 자동 재시도하는 중",
                "별도 조치 없이 잠시 기다립니다. 시스템이 자동으로 다음 주기에 다시 처리합니다.",
            ],
            [
                "재시도가 한도를 넘어 강제로 종료됨",
                "카드가 완료 컬럼으로 이동했으나 로그에 \"[Error]\"로 시작하는 메시지 표시",
                "여러 번의 자동 처리 시도에도 같은 오류가 반복된 상태",
                "카드 상세의 오류 메시지를 확인하고 운영 담당자에게 문의합니다. 입력 데이터나 프로세스 정의를 보완해야 할 수 있습니다.",
            ],
            [
                "외부 고객에게 안내 메일이 가지 않음",
                "다음 단계 카드는 만들어졌으나 고객이 이메일을 받지 못함",
                "프로세스 정보에 등록된 고객 이메일이 비어 있거나 잘못 입력된 경우",
                "프로세스 정보의 고객 이메일을 확인하고 운영 담당자에게 수정 요청합니다.",
            ],
        ],
        [3.5, 4.0, 4.5, 5.0],
    )

    # 12. 권한 및 역할별 기능 차이
    add_heading(doc, "12. 권한 및 역할별 기능 차이")
    add_table(
        doc,
        ["역할", "접근 가능한 메뉴", "수행 가능한 작업", "제한 사항"],
        [
            ["일반 사용자", "할 일 보드, 본인 카드 상세", "본인 카드 입력·제출, 결과·로그 확인", "타인의 카드 수정 불가, 인스턴스 전체 관리 불가"],
            ["프로세스 담당자", "할 일 보드, 프로세스 인스턴스 목록 및 상세", "인스턴스 상태 모니터링, 정체 카드 확인", "프로세스 정의 자체 수정은 별도 권한 필요"],
            ["외부 고객", "안내 이메일에 포함된 외부 폼 링크", "본인 작업 입력 및 제출", "본 화면 로그인 불필요, 보드/인스턴스 접근 불가"],
        ],
        [3.2, 4.5, 5.5, 4.5],
    )

    # 13. FAQ
    add_heading(doc, "13. FAQ / 자주 묻는 질문")
    add_table(
        doc,
        ["질문", "답변"],
        [
            ["카드가 제출 대기에서 자동으로 진행되지 않습니다.", "시스템은 일정 주기로 새 카드를 감지해 처리합니다. 잠시 기다린 후 새로고침해 보세요. 오랫동안 변하지 않으면 정체된 처리가 자동으로 해제되고 재처리됩니다."],
            ["카드 상태가 \"대기\"로 바뀌었습니다. 무엇을 해야 하나요?", "카드 상세의 사유 메시지를 확인하세요. 추가 입력이 필요한 경우 안내가 표시되며, 그렇지 않으면 운영 담당자에게 문의합니다."],
            ["다음 단계 카드가 보이지 않습니다.", "본인이 다음 단계의 담당자가 아닐 수 있습니다. 인스턴스 상세에서 현재 진행 단계와 담당자를 확인하세요."],
            ["로그에 \"[Error]\" 메시지가 나타났습니다.", "자동 재시도가 한도를 넘어 강제로 종료된 상태입니다. 운영 담당자에게 문의해 원인을 확인하고 후속 조치를 받으세요."],
            ["외부 고객에게 안내 이메일이 도착하지 않습니다.", "프로세스 정보의 외부 고객 이메일이 정확히 등록되어 있는지 확인하세요. 등록은 되어 있지만 수신되지 않는 경우 운영 담당자에게 문의합니다."],
            ["내가 한 작업을 다시 확인하고 싶습니다.", "완료 컬럼의 카드를 클릭해 상세 화면을 열면 입력 내용, 결과, 처리 로그를 다시 확인할 수 있습니다."],
        ],
        [7.0, 11.0],
    )

    # 14. 효과적인 사용 팁
    add_heading(doc, "14. 효과적인 사용 팁")
    add_bullets(
        doc,
        [
            "카드를 제출하기 전에 필수 입력 항목과 첨부 내용이 모두 갖추어졌는지 한 번 더 확인합니다.",
            "자동 진행이 진행 중일 때는 카드 상태가 곧 바뀔 수 있으므로, 여러 번 같은 작업을 반복 제출하지 않습니다.",
            "카드가 \"대기\" 상태가 되면 먼저 카드 상세의 사유 메시지를 확인한 뒤 추가 조치 여부를 결정합니다.",
            "오류 메시지(\"[Error]\")가 보이면 임의로 다시 제출하지 말고, 운영 담당자에게 문의해 원인 확인 후 조치를 받습니다.",
            "외부 고객 단계가 포함된 프로세스는, 시작 전에 고객 이메일 주소가 올바른지 확인해 두면 안내 메일 발송 실패를 예방할 수 있습니다.",
        ],
    )

    return doc


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
