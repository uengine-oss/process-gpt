"""프로세스 정의 검색 사용자 매뉴얼 생성 스크립트.

OpenSpec ``completion_process-definition-search`` 사양과 E2E 산출물을 근거로
최종 사용자가 읽을 수 있는 DOCX 매뉴얼을 생성합니다. 시각 스타일은
``.claude/skills/docx-user-manual/STYLE_REFERENCE.py`` 를 기준으로 합니다.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SPEC_DIR = Path(__file__).resolve().parent.parent
OUTPUT_PATH = SPEC_DIR / "docs" / "completion_process-definition-search-user-manual.docx"
SCREENSHOT_DIR = SPEC_DIR / "e2e" / "results" / "screenshots"

DOC_VERSION = "v1.0"
DOC_DATE = date(2026, 5, 28)
ORG_NAME = "Process-GPT"
SERVICE_NAME = "Process-GPT"
DOC_TITLE = "프로세스 정의 검색 사용자 매뉴얼"
FONT_NAME = "Malgun Gothic"
PRIMARY_COLOR = "1F4E79"
CAPTION_COLOR = "666666"
TABLE_HEADER_FILL = "D9EAF7"


SCREENSHOTS = {
    "initial": "process-gpt-completion_process-definition-search-01-search-initial.png",
    "input": "process-gpt-completion_process-definition-search-01-search-input.png",
    "result": "process-gpt-completion_process-definition-search-01-search-result.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False, align_center: bool = True) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차는 Word에서 필드 업데이트 시 페이지 번호와 함께 표시됩니다."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, bold=True, color=PRIMARY_COLOR)


def add_paragraph(doc: Document, text: str = "", bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        run = paragraph.add_run(f"{idx}.\t{item}")
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, TABLE_HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, align_center=False)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_image(doc: Document, key: str, caption: str) -> None:
    path = SCREENSHOT_DIR / SCREENSHOTS[key]
    if not path.exists():
        add_paragraph(doc, f"[화면 예시 누락] {path.name}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.0))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=CAPTION_COLOR)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SERVICE_NAME)
    set_run_font(run, size=18, bold=True, color=PRIMARY_COLOR)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DOC_TITLE)
    set_run_font(run, size=28, bold=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 버전", DOC_VERSION],
            ["작성일", DOC_DATE.isoformat()],
            ["작성자/조직", ORG_NAME],
            ["문서 구분", "최종 사용자 매뉴얼"],
            ["대상 기능", "프로세스 정의 검색"],
        ],
        [4.0, 12.0],
    )
    doc.add_page_break()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{DOC_TITLE} | ")
    set_run_font(run, size=9, color=CAPTION_COLOR)
    add_page_number(footer)

    doc.core_properties.title = DOC_TITLE
    doc.core_properties.subject = "프로세스 정의 검색"
    doc.core_properties.author = ORG_NAME
    enable_field_update_on_open(doc)


def add_scenario(
    doc: Document,
    title: str,
    purpose: str,
    preconditions: list[str],
    steps: list[str],
    expected: list[str],
    screenshot_key: str | None = None,
    caption: str | None = None,
) -> None:
    add_heading(doc, title, level=2)
    add_paragraph(doc, "이럴 때 사용합니다", bold=True)
    add_paragraph(doc, purpose)
    add_paragraph(doc, "사용 전 확인", bold=True)
    add_bullets(doc, preconditions)
    add_paragraph(doc, "따라 하기", bold=True)
    add_numbers(doc, steps)
    add_paragraph(doc, "화면에서 확인할 내용", bold=True)
    add_bullets(doc, expected)
    if screenshot_key and caption:
        add_image(doc, screenshot_key, caption)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    # 1. 문서 정보 및 변경 이력
    add_heading(doc, "1. 문서 정보 및 변경 이력")
    add_table(
        doc,
        ["문서 버전", "변경일", "변경 내용", "작성자", "검토자"],
        [[DOC_VERSION, DOC_DATE.isoformat(), "초안 작성", ORG_NAME, "서비스 담당자"]],
        [2.6, 3.0, 8.2, 3.0, 3.0],
    )

    # 2. 목차
    add_heading(doc, "2. 목차")
    add_paragraph(doc, "아래 목차는 Word에서 문서를 열 때 필드를 업데이트하면 페이지 번호가 반영됩니다.")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    # 3. 개요
    add_heading(doc, "3. 개요")
    add_paragraph(
        doc,
        "본 매뉴얼은 Process-GPT 화면 상단의 검색바에서 자연어로 프로세스 정의를 찾아보는 방법을 안내합니다.",
    )
    add_bullets(
        doc,
        [
            "서비스 목적: 사용자가 자연어 질의로 시작하거나 참고할 프로세스 정의를 의미 기반으로 빠르게 찾도록 돕습니다.",
            "매뉴얼 목적: 화면 진입, 검색어 입력, 결과 확인, 결과가 없을 때의 대응까지 따라 하기 형식으로 안내합니다.",
            "기능 범위: 헤더 검색바를 통한 유사 프로세스 정의 검색, 결과 패널 활용, 테넌트별 결과 분리 동작을 다룹니다.",
        ],
    )

    # 4. 대상 사용자
    add_heading(doc, "4. 대상 사용자")
    add_table(
        doc,
        ["사용자 역할", "주요 사용 목적", "권장 사용 범위"],
        [
            ["일반 사용자", "자신이 시작하려는 업무에 적합한 기존 프로세스 정의를 찾습니다.", "헤더 검색바 자연어 검색"],
            ["업무 담당자", "참고 가능한 유사 프로세스 정의를 빠르게 확인하고 활용합니다.", "검색 결과 항목 열람"],
            ["검토자", "테넌트(소속 조직) 안에서만 결과가 표시되는지 확인합니다.", "검색 결과 확인"],
        ],
        [3.5, 8.0, 5.0],
    )

    # 5. 사용 전 확인 사항
    add_heading(doc, "5. 사용 전 확인 사항")
    add_table(
        doc,
        ["확인 항목", "내용"],
        [
            ["계정", "본인 소속 조직(테넌트)으로 로그인할 수 있는 계정이 있습니다."],
            ["접속 URL", "조직에서 안내받은 Process-GPT 접속 주소를 사용합니다."],
            ["브라우저", "최신 버전의 Chrome 또는 Edge 사용을 권장합니다."],
            ["검색어 준비", "찾고자 하는 업무를 떠올릴 수 있는 자연어 키워드(예: 휴가, 구매, 출장)를 준비합니다."],
            ["프로세스 정의 등록", "소속 조직에 검색 대상 프로세스 정의가 한 건 이상 등록되어 있어야 합니다."],
        ],
        [4.0, 12.0],
    )

    # 6. 시작하기
    add_heading(doc, "6. 시작하기")
    add_paragraph(doc, "처음 사용하는 경우 다음 절차로 화면에 진입한 뒤 검색을 수행합니다.")
    add_numbers(
        doc,
        [
            "브라우저에서 조직에서 안내받은 Process-GPT 주소로 이동합니다.",
            "로그인 화면에서 이메일과 비밀번호를 입력한 뒤 로그인 버튼을 누릅니다.",
            "로그인이 완료되면 화면 상단의 헤더 영역에 검색바(돋보기 아이콘 옆 입력창)가 표시되는지 확인합니다.",
            "검색바에 자연어 키워드를 입력하고 Enter 키를 눌러 결과 패널을 엽니다.",
        ],
    )
    add_image(doc, "initial", "그림 1. 로그인 직후 헤더 검색바 초기 상태")

    # 7. 화면 구성
    add_heading(doc, "7. 화면 구성")
    add_paragraph(doc, "프로세스 정의 검색과 관련된 주요 화면 영역은 다음과 같습니다.")
    add_table(
        doc,
        ["영역", "위치", "역할"],
        [
            ["헤더 검색바", "화면 상단 가운데/우측 영역", "자연어 키워드를 입력해 검색을 시작합니다."],
            ["돋보기 아이콘", "검색바 왼쪽", "검색 입력 영역임을 나타냅니다."],
            ["검색 결과 패널", "검색바 아래로 펼쳐지는 영역", "검색어 입력 후 결과 카테고리와 항목을 표시합니다."],
            ["유사한 프로세스 정의 카테고리", "검색 결과 패널 내", "입력한 키워드와 의미가 유사한 프로세스 정의 목록을 보여줍니다."],
            ["결과 항목", "유사한 프로세스 정의 카테고리 내", "선택 가능한 링크 형태의 프로세스 정의 제목입니다."],
        ],
        [4.5, 5.5, 7.0],
    )

    # 8. 주요 사용 흐름
    add_heading(doc, "8. 주요 사용 흐름")

    add_scenario(
        doc,
        "8.1 헤더 검색바로 유사 프로세스 정의 찾기",
        "시작하려는 업무가 떠오르지만 정확한 프로세스 정의 이름을 모를 때, 자연어 키워드만으로 유사 프로세스 정의 후보를 빠르게 확인할 수 있습니다.",
        [
            "Process-GPT에 로그인되어 헤더 검색바가 보입니다.",
            "소속 조직에 검색 대상이 될 만한 프로세스 정의가 등록되어 있습니다.",
            "찾으려는 업무를 표현할 수 있는 자연어 키워드(예: 휴가)를 떠올린 상태입니다.",
        ],
        [
            "헤더 영역의 돋보기 아이콘 옆 검색바를 클릭합니다.",
            "찾고자 하는 업무를 표현하는 자연어 키워드(예: 휴가)를 입력합니다.",
            "Enter 키를 눌러 검색을 실행합니다.",
            "검색바 아래로 펼쳐지는 결과 패널에서 ‘유사한 프로세스 정의’ 카테고리가 나타날 때까지 잠시 기다립니다.",
            "결과 항목 중 원하는 프로세스 정의 제목을 확인하고, 필요 시 항목을 클릭해 다음 작업으로 이동합니다.",
        ],
        [
            "검색 결과 패널이 펼쳐지고 ‘유사한 프로세스 정의’ 섹션 헤더가 보입니다.",
            "섹션 안에 입력한 키워드와 의미가 유사한 프로세스 정의 제목(예: 휴가신청)이 표시됩니다.",
            "결과 항목은 클릭 가능한 링크 형태로 렌더링됩니다.",
            "최대 3건까지 유사 후보가 나타날 수 있습니다.",
        ],
        "input",
        "그림 2. 검색바에 자연어 키워드를 입력한 상태",
    )

    add_image(doc, "result", "그림 3. ‘유사한 프로세스 정의’ 카테고리에 결과가 표시된 화면")

    add_scenario(
        doc,
        "8.2 검색 결과가 없을 때 다시 시도하기",
        "입력한 키워드와 의미적으로 유사한 프로세스 정의가 조직 내에 등록되어 있지 않을 때, 사용자가 빈 결과를 인지하고 키워드를 바꿔 다시 시도하는 흐름입니다.",
        [
            "헤더 검색바에서 한 번 이상 검색을 시도했습니다.",
            "현재 조직에 유사 후보로 보일 만한 프로세스 정의가 등록되지 않았을 수 있습니다.",
        ],
        [
            "헤더 검색바를 다시 클릭합니다.",
            "기존 키워드를 지우고 더 일반적인 자연어 표현(예: 휴가 → 연차, 휴가신청)으로 바꿔 입력합니다.",
            "Enter 키를 눌러 다시 검색합니다.",
            "‘유사한 프로세스 정의’ 카테고리에 결과가 표시되는지 확인합니다.",
            "여전히 결과가 비어 있다면 ‘효과적인 질문 작성 팁’ 섹션을 참고해 다른 표현으로 시도합니다.",
        ],
        [
            "검색은 정상적으로 종료되며 오류 메시지 없이 빈 결과 상태가 표시됩니다.",
            "‘유사한 프로세스 정의’ 카테고리에 항목이 없을 수 있으며, 이는 조직 내에 유사 후보가 없는 정상 상태입니다.",
            "다른 키워드로 다시 검색하면 새로운 결과 목록이 표시될 수 있습니다.",
        ],
    )

    add_scenario(
        doc,
        "8.3 소속 조직 안에서만 결과가 표시되는지 확인하기",
        "프로세스 정의 검색은 사용자가 접속한 조직(테넌트) 안에 등록된 프로세스 정의만 대상으로 합니다. 본 흐름은 다른 조직의 프로세스 정의가 결과에 섞여 나오지 않는지 사용자가 직접 확인하는 절차입니다.",
        [
            "본인 계정이 소속된 조직으로 로그인되어 있습니다.",
            "다른 조직에만 존재한다고 알려진 키워드가 있다면 미리 확인합니다.",
        ],
        [
            "헤더 검색바에 본인 조직에서 사용할 법한 키워드(예: 휴가)를 입력하고 검색합니다.",
            "‘유사한 프로세스 정의’ 카테고리에 본인 조직의 프로세스 정의 제목이 표시되는지 확인합니다.",
            "다른 조직에만 존재하는 키워드(예: 회의실)를 입력해 다시 검색합니다.",
            "본인 조직에 해당 프로세스 정의가 없다면 결과가 비어 있는지 확인합니다.",
        ],
        [
            "검색 결과는 본인이 로그인한 조직(테넌트)의 프로세스 정의로 한정되어 표시됩니다.",
            "다른 조직에서만 등록된 프로세스 정의 제목은 결과 목록에 나타나지 않습니다.",
            "결과가 비어 있는 경우에도 오류 없이 정상 종료됩니다.",
        ],
    )

    # 9. 화면 항목 및 옵션 설명
    add_heading(doc, "9. 화면 항목 및 옵션 설명")
    add_table(
        doc,
        ["화면 항목", "사용 여부", "형식/예시", "설명"],
        [
            ["헤더 검색바 입력창", "필수", "자연어 텍스트 (예: 휴가, 구매요청, 출장)", "찾고자 하는 업무를 자연어 키워드로 입력합니다."],
            ["Enter 키", "필수", "키보드 입력", "입력한 키워드로 검색을 실행합니다."],
            ["검색 결과 패널", "자동 표시", "검색바 아래 영역", "검색 실행 후 결과 카테고리와 항목을 보여줍니다."],
            ["유사한 프로세스 정의 카테고리", "자동 표시", "결과 패널 내 섹션", "의미가 유사한 프로세스 정의 후보를 묶어 보여줍니다."],
            ["결과 항목 링크", "선택", "프로세스 정의 제목 텍스트", "클릭하면 해당 프로세스 정의를 활용하는 다음 단계로 이동합니다."],
        ],
        [4.5, 2.5, 4.5, 5.5],
    )

    # 10. 결과 확인
    add_heading(doc, "10. 결과 확인")
    add_paragraph(doc, "검색을 실행한 뒤 다음 항목을 확인합니다.")
    add_bullets(
        doc,
        [
            "‘유사한 프로세스 정의’ 카테고리가 결과 패널에 표시되는지 확인합니다.",
            "카테고리 안에 표시된 프로세스 정의 제목이 입력한 키워드와 의미적으로 가까운지 확인합니다.",
            "최대 3건의 유사 후보가 나타날 수 있으며, 결과가 1~2건이거나 비어 있을 수도 있습니다.",
            "원하는 항목을 클릭해 다음 작업으로 이어갈 수 있는지 확인합니다.",
            "다른 키워드로 다시 검색해 결과 비교가 필요한 경우 검색바를 다시 사용합니다.",
        ],
    )

    # 11. 오류 및 예외 상황
    add_heading(doc, "11. 오류 및 예외 상황")
    add_table(
        doc,
        ["증상", "예상 원인", "사용자 조치"],
        [
            [
                "검색바가 헤더에 보이지 않습니다.",
                "로그인이 풀렸거나 화면 진입 권한이 없습니다.",
                "로그아웃 후 다시 로그인하고, 계속 보이지 않으면 관리자에게 문의합니다.",
            ],
            [
                "검색 결과 패널이 열리지 않거나 빈 상태로 유지됩니다.",
                "조직에 유사한 프로세스 정의가 등록되어 있지 않거나, 입력한 키워드가 너무 구체적입니다.",
                "키워드를 더 일반적인 표현으로 바꾸거나 동의어를 사용해 다시 검색합니다.",
            ],
            [
                "‘유사한 프로세스 정의’ 카테고리가 비어 있습니다.",
                "조직(테넌트) 안에 유사 후보가 없거나, 임시 검색 오류로 빈 결과가 반환됩니다.",
                "잠시 후 다른 키워드로 다시 시도하고, 반복되면 관리자에게 문의합니다.",
            ],
            [
                "다른 조직에서 본 적 있는 프로세스 정의가 결과에 보이지 않습니다.",
                "검색 범위가 현재 로그인한 조직으로 한정되어 있습니다.",
                "본인 조직 기준의 키워드로 다시 검색하거나, 필요한 경우 해당 조직 담당자에게 문의합니다.",
            ],
        ],
        [4.0, 5.5, 6.5],
    )

    # 12. 권한 및 역할별 기능 차이
    add_heading(doc, "12. 권한 및 역할별 기능 차이")
    add_table(
        doc,
        ["사용자 역할", "검색 사용", "결과 열람"],
        [
            ["일반 사용자", "가능", "본인이 속한 조직의 프로세스 정의만 표시됩니다."],
            ["업무 담당자", "가능", "본인이 속한 조직의 프로세스 정의만 표시됩니다."],
            ["검토자", "가능", "본인이 속한 조직의 프로세스 정의만 표시됩니다."],
            ["미로그인 사용자", "권한 필요", "헤더 검색바 사용을 위해 먼저 로그인해야 합니다."],
        ],
        [4.5, 3.0, 8.5],
    )

    # 13. FAQ
    add_heading(doc, "13. FAQ / 자주 묻는 질문")
    add_table(
        doc,
        ["질문", "답변"],
        [
            [
                "검색어를 다 입력했는데도 결과가 보이지 않습니다.",
                "Enter 키를 눌렀는지 확인하고, 키워드를 더 짧고 일반적인 단어로 바꿔 다시 검색해 보세요.",
            ],
            [
                "‘유사한 프로세스 정의’가 정확히 무엇을 보여주는 건가요?",
                "입력한 자연어 키워드와 의미적으로 가까운 프로세스 정의 제목 후보를 최대 3건까지 보여줍니다.",
            ],
            [
                "다른 회사/조직의 프로세스 정의도 함께 보고 싶습니다.",
                "검색 결과는 현재 로그인한 조직 안의 프로세스 정의로만 한정됩니다. 다른 조직 정보를 보려면 해당 조직 담당자에게 문의하세요.",
            ],
            [
                "결과가 1건뿐일 때는 어떻게 해야 하나요?",
                "후보가 1건만 존재하거나 의미가 가까운 후보가 1건만 있는 정상 상태일 수 있습니다. 필요하면 키워드를 바꿔 다른 후보가 있는지 다시 검색해 보세요.",
            ],
            [
                "결과가 전혀 없는데 오류 메시지도 보이지 않습니다.",
                "조직에 유사 후보가 없는 정상 상태입니다. 다른 키워드로 다시 시도하거나 ‘효과적인 질문 작성 팁’을 참고하세요.",
            ],
        ],
        [6.5, 11.0],
    )

    # 14. 효과적인 질문 작성 팁
    add_heading(doc, "14. 효과적인 질문 작성 팁")
    add_paragraph(doc, "검색 결과를 더 잘 받기 위해 다음 방법을 권장합니다.")
    add_bullets(
        doc,
        [
            "업무명을 그대로 떠올려 짧은 자연어로 입력합니다. 예: 휴가, 구매, 출장, 회의실.",
            "정확한 프로세스 정의 이름을 몰라도 됩니다. 비슷한 의미의 단어로 시도하세요. 예: 휴가 → 연차, 휴가신청.",
            "결과가 없으면 더 일반적인 표현으로 바꿔 다시 시도합니다. 예: 출장신청서 → 출장.",
            "결과가 너무 많거나 모호하면 두 단어를 함께 사용해 의미를 좁힙니다. 예: 휴가 신청.",
            "검색은 본인 조직 범위에서만 동작하므로, 다른 조직에 있는 프로세스 정의는 결과에 표시되지 않습니다.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
