"""DOCX user manual generator for completion_process-activity-rework.

Generates an end-user manual that walks process operators through identifying
completed activities, choosing the rework scope, and starting rework from the
WorkItem detail screen. Style mirrors STYLE_REFERENCE.py.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SPEC_ROOT = HERE.parent
OUTPUT_PATH = HERE / "completion_process-activity-rework-user-manual.docx"
SCREENSHOT_DIR = SPEC_ROOT / "e2e" / "results" / "screenshots"

DOC_VERSION = "v1.0"
DOC_DATE = date(2026, 5, 27)
ORG_NAME = "Process-GPT"
SERVICE_NAME = "Process-GPT 업무 자동화"
DOC_TITLE = "프로세스 활동 재작업 사용자 매뉴얼"
FONT_NAME = "Malgun Gothic"
PRIMARY_COLOR = "1F4E79"
CAPTION_COLOR = "666666"
TABLE_HEADER_FILL = "D9EAF7"


SCREENSHOTS = {
    "workitem-detail-done": "process-gpt-completion_process-activity-rework-01-workitem-detail-done.png",
    "rework-dialog-open": "process-gpt-completion_process-activity-rework-01-rework-dialog-open.png",
    "rework-dialog-include-all": "process-gpt-completion_process-activity-rework-01-rework-dialog-include-all.png",
    "rework-dialog-submit-ready": "process-gpt-completion_process-activity-rework-02-rework-dialog-submit-ready.png",
    "instance-after-rework": "process-gpt-completion_process-activity-rework-02-instance-after-rework.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차는 Word에서 필드 업데이트 시 페이지 번호와 함께 표시됩니다."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, bold=True, color=PRIMARY_COLOR)


def add_paragraph(doc: Document, text: str = "", bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold if bold else None)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        run = paragraph.add_run(f"{idx}.\t{item}")
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, TABLE_HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_image(doc: Document, key: str, caption: str) -> None:
    path = SCREENSHOT_DIR / SCREENSHOTS[key]
    if not path.exists():
        add_paragraph(doc, f"[화면 예시 누락] {path.name}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.0))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=CAPTION_COLOR)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SERVICE_NAME)
    set_run_font(run, size=18, bold=True, color=PRIMARY_COLOR)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DOC_TITLE)
    set_run_font(run, size=28, bold=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 버전", DOC_VERSION],
            ["작성일", DOC_DATE.isoformat()],
            ["작성자/조직", ORG_NAME],
            ["문서 구분", "사용자 참조 매뉴얼"],
            ["대상 서비스", SERVICE_NAME],
        ],
        [4.0, 12.0],
    )
    doc.add_page_break()


def add_scenario(
    doc: Document,
    title: str,
    purpose: str,
    preconditions: list[str],
    steps: list[str],
    expected: list[str],
    images: list[tuple[str, str]] | None = None,
) -> None:
    add_heading(doc, title, level=2)
    add_paragraph(doc, f"이럴 때 사용합니다: {purpose}")
    add_paragraph(doc, "사용 전 확인", bold=True)
    add_bullets(doc, preconditions)
    add_paragraph(doc, "따라 하기", bold=True)
    add_numbers(doc, steps)
    add_paragraph(doc, "화면에서 확인할 내용", bold=True)
    add_bullets(doc, expected)
    if images:
        for key, caption in images:
            add_image(doc, key, caption)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{SERVICE_NAME} 사용자 매뉴얼 | ")
    set_run_font(run, size=9, color=CAPTION_COLOR)
    add_page_number(footer)

    doc.core_properties.title = DOC_TITLE
    doc.core_properties.subject = SERVICE_NAME
    doc.core_properties.author = ORG_NAME
    enable_field_update_on_open(doc)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. 문서 정보 및 변경 이력")
    add_table(
        doc,
        ["문서 버전", "변경일", "변경 내용", "작성자", "검토자"],
        [[DOC_VERSION, DOC_DATE.isoformat(), "초안 작성", ORG_NAME, "서비스 담당자"]],
        [2.6, 3.0, 8.2, 3.0, 3.0],
    )

    add_heading(doc, "2. 목차")
    add_paragraph(doc, "아래 목차는 Word에서 문서를 열 때 필드를 업데이트하면 페이지 번호가 반영됩니다.")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    add_heading(doc, "3. 개요")
    add_paragraph(
        doc,
        "이 매뉴얼은 이미 완료된 프로세스 단계를 다시 수행해야 할 때, 사용자가 화면에서 재작업 범위를 선택하고 안전하게 재작업을 시작하는 방법을 안내합니다.",
    )
    add_bullets(
        doc,
        [
            "서비스 목적: 완료된 업무 단계를 사용자가 직접 선택해 다시 수행할 수 있도록 지원합니다.",
            "매뉴얼 목적: 완료된 작업 상세 화면 진입, 재작업 다이얼로그 사용, 재작업 시작 및 결과 확인까지 순서대로 안내합니다.",
            "기능 범위: 재작업 후보 활동 조회, 재작업 범위 선택, 재작업 시작, 인스턴스 화면에서의 결과 확인을 포함합니다.",
        ],
    )

    add_heading(doc, "4. 대상 사용자")
    add_table(
        doc,
        ["사용자 역할", "주요 사용 목적", "권장 사용 범위"],
        [
            ["프로세스 담당자", "자신이 완료한 작업을 다시 수행해야 할 때 재작업을 시작합니다.", "본인 소유 작업 재작업"],
            ["업무 운영자", "완료된 인스턴스에서 잘못된 결과를 바로잡기 위해 이후 단계까지 다시 수행합니다.", "후속 단계 포함 재작업"],
            ["검토자", "재작업 결과 인스턴스 화면에서 현재 진행 상태를 확인합니다.", "인스턴스 상태 확인"],
        ],
        [4.0, 8.0, 5.0],
    )

    add_heading(doc, "5. 사용 전 확인 사항")
    add_bullets(
        doc,
        [
            "Process-GPT에 접속할 수 있는 사용자 계정과 비밀번호가 있습니다.",
            "재작업하려는 작업이 본인 소유이며 \"완료\" 상태인지 확인합니다.",
            "재작업 대상 인스턴스가 어떤 프로세스인지, 어떤 단계까지 다시 수행할지 사전에 결정해 둡니다.",
            "권장 브라우저(Chrome 등 최신 버전)를 사용합니다.",
        ],
    )

    add_heading(doc, "6. 시작하기")
    add_numbers(
        doc,
        [
            "브라우저 주소창에 서비스 접속 주소를 입력해 로그인 화면을 엽니다.",
            "본인 계정으로 로그인합니다.",
            "할 일 목록(Todolist)에서 재작업할 완료 작업을 찾아 상세 화면으로 이동합니다.",
            "상세 화면 상단에 \"재작업\" 버튼이 보이는지 확인합니다.",
        ],
    )
    add_image(doc, "workitem-detail-done", "그림 1. 완료한 작업 상세 화면에서 재작업 버튼이 노출됩니다.")

    add_heading(doc, "7. 화면 구성")
    add_table(
        doc,
        ["화면 영역", "역할", "주요 내용"],
        [
            ["작업 상세 화면", "완료된 작업의 정보 확인 및 재작업 진입점 제공", "작업명, 상태, 재작업 버튼"],
            ["재작업 다이얼로그", "어느 범위까지 다시 수행할지 선택", "범위 라디오 옵션, 대상 활동 칩, 시작 버튼"],
            ["활동 칩 영역", "선택한 범위에 따라 다시 수행될 활동 표시", "강조 표시된 활동 칩"],
            ["인스턴스 상세 화면", "재작업 시작 후 진행 상태 확인", "현재 진행 단계, 다음 단계 목록"],
        ],
        [4.5, 5.5, 7.0],
    )

    add_heading(doc, "8. 주요 사용 흐름")

    add_scenario(
        doc,
        "8.1 재작업할 범위를 다이얼로그에서 확인하기",
        "완료한 작업을 다시 수행해야 할 때, 어떤 단계까지 다시 진행될지 선택지를 미리 확인합니다.",
        [
            "완료된 본인 작업의 상세 화면이 열려 있습니다.",
            "다시 수행할 단계 범위에 대한 결정이 준비되어 있습니다.",
        ],
        [
            "작업 상세 화면에서 \"재작업\" 버튼을 클릭합니다.",
            "재작업 다이얼로그가 열리고 \"현재 활동만\", \"참조 활동 포함\", \"이후 모든 활동 포함\" 세 가지 옵션이 표시되는지 확인합니다.",
            "원하는 범위 옵션을 클릭해 대상 활동 칩이 강조되는지 확인합니다.",
            "범위를 다시 검토하려면 다른 옵션을 선택해 칩 구성이 어떻게 바뀌는지 비교합니다.",
            "선택을 취소하려면 다이얼로그 우측 상단 닫기(X) 버튼을 누릅니다.",
        ],
        [
            "재작업 버튼을 누르면 다이얼로그가 화면 중앙에 표시됩니다.",
            "세 가지 범위 옵션과 대상 활동 칩이 모두 보입니다.",
            "범위 옵션을 선택하면 해당 범위에 포함되는 활동 칩이 강조됩니다.",
            "닫기를 누르면 다이얼로그가 사라지고 작업 상세 화면이 그대로 유지됩니다.",
        ],
        [
            ("rework-dialog-open", "그림 2. 어디까지 재작업할지 선택하는 다이얼로그가 표시됩니다."),
            ("rework-dialog-include-all", "그림 3. 이후 활동까지 모두 재작업 대상으로 선택한 화면입니다."),
        ],
    )

    add_scenario(
        doc,
        "8.2 이후 단계까지 모두 다시 수행하기",
        "현재 단계뿐 아니라 이후 단계까지 결과를 다시 만들어야 할 때 사용합니다.",
        [
            "재작업 다이얼로그가 열려 있습니다.",
            "이후 단계까지 다시 수행해도 되는지 업무 관점에서 결정되어 있습니다.",
        ],
        [
            "다이얼로그에서 \"이후 모든 활동 포함\" 옵션을 선택합니다.",
            "강조 표시된 활동 칩이 다시 수행될 단계 전체를 포함하는지 확인합니다.",
            "다이얼로그 하단의 \"재작업 시작\" 버튼이 활성화되었는지 확인합니다.",
            "\"재작업 시작\" 버튼을 클릭합니다.",
            "화면이 인스턴스 상세 화면으로 이동할 때까지 잠시 기다립니다.",
        ],
        [
            "선택한 범위에 포함된 활동 칩이 강조되어 표시됩니다.",
            "재작업 시작 버튼이 활성화 상태로 바뀝니다.",
            "버튼을 누르면 인스턴스 상세 화면으로 자동 이동합니다.",
            "이동한 화면에서 시작 활동은 \"진행 중\" 상태, 이후 활동은 \"대기\" 상태로 표시됩니다.",
        ],
        [
            ("rework-dialog-submit-ready", "그림 4. 재작업 범위를 선택하고 시작 버튼이 활성화된 다이얼로그입니다."),
            ("instance-after-rework", "그림 5. 재작업이 시작되어 인스턴스 화면으로 이동했습니다."),
        ],
    )

    add_heading(doc, "9. 화면 항목 및 옵션 설명")
    add_table(
        doc,
        ["화면 항목/옵션", "사용 여부", "기본값", "설명"],
        [
            ["재작업 버튼", "조건부 노출", "-", "본인이 완료한 작업 상세 화면에서만 표시됩니다."],
            ["현재 활동만", "선택", "기본 선택", "지금 보고 있는 단계만 다시 수행합니다."],
            ["참조 활동 포함", "선택", "-", "현재 단계와 함께, 다시 수행에 필요한 참조 단계까지 함께 다시 수행합니다."],
            ["이후 모든 활동 포함", "선택", "-", "현재 단계와 이후에 이어지는 모든 단계를 다시 수행합니다."],
            ["대상 활동 칩", "자동 표시", "-", "선택한 범위에 포함되는 활동들을 강조 표시합니다."],
            ["재작업 시작 버튼", "필수", "-", "선택한 범위로 재작업을 실제로 시작합니다."],
            ["닫기(X)", "선택", "-", "다이얼로그를 닫고 작업 상세 화면으로 돌아갑니다."],
        ],
        [4.0, 3.0, 3.0, 7.0],
    )

    add_heading(doc, "10. 결과 확인")
    add_bullets(
        doc,
        [
            "재작업이 시작되면 화면이 자동으로 인스턴스 상세 화면으로 이동합니다.",
            "재작업의 시작 단계는 \"진행 중\" 상태로 표시됩니다.",
            "함께 다시 수행하기로 선택한 이후 단계들은 \"대기\" 상태로 표시됩니다.",
            "기존 완료 이력은 인스턴스에 유지되며, 재작업으로 만들어진 새 단계는 별도의 새 작업으로 추가됩니다.",
            "재작업 단계가 다시 완료되면 평소와 동일하게 \"완료\" 상태로 전환됩니다.",
        ],
    )

    add_heading(doc, "11. 오류 및 예외 상황")
    add_table(
        doc,
        ["증상", "원인", "사용자 조치"],
        [
            [
                "재작업 버튼이 보이지 않습니다.",
                "현재 작업이 완료 상태가 아니거나 본인 소유가 아닙니다.",
                "할 일 목록에서 본인이 완료한 작업의 상세 화면인지 다시 확인합니다.",
            ],
            [
                "재작업 다이얼로그에 활동 칩이 표시되지 않습니다.",
                "재작업 후보 활동을 불러오는 중이거나 후속 단계가 없는 마지막 단계입니다.",
                "잠시 후 다시 확인하고, 마지막 단계인 경우 \"현재 활동만\" 옵션을 사용합니다.",
            ],
            [
                "재작업 시작 버튼을 눌렀는데 화면이 이동하지 않습니다.",
                "네트워크 지연 또는 일시적인 오류로 응답이 늦어지는 상황입니다.",
                "잠시 기다린 후에도 이동하지 않으면 페이지를 새로 고치고 다시 시도합니다.",
            ],
            [
                "다른 사용자가 동일 인스턴스에서 작업 중이어서 재작업이 거절됩니다.",
                "동일 인스턴스에 진행 중인 다른 단계가 있어 안전하게 되돌릴 수 없습니다.",
                "해당 단계가 마무리된 뒤 다시 재작업을 시도하거나 운영 담당자에게 문의합니다.",
            ],
            [
                "오류 메시지로 작업을 찾을 수 없다는 안내가 표시됩니다.",
                "대상 인스턴스 또는 활동이 더 이상 존재하지 않습니다.",
                "할 일 목록에서 인스턴스가 유지되고 있는지 확인한 뒤 다시 진입합니다.",
            ],
        ],
        [5.0, 5.5, 6.5],
    )

    add_heading(doc, "12. 권한 및 역할별 기능 차이")
    add_table(
        doc,
        ["역할", "재작업 버튼 사용", "재작업 범위 선택", "비고"],
        [
            ["본인 작업 담당자", "가능", "현재/참조/이후 모두 사용", "본인이 완료한 작업에서만 노출됩니다."],
            ["같은 인스턴스의 다른 사용자", "권한 필요", "권한 필요", "본인 소유 작업이 아닌 경우 버튼이 보이지 않을 수 있습니다."],
            ["운영자/관리자", "권한 필요", "권한 필요", "운영자 권한 정책에 따라 결정되며, 정책이 명확하지 않으면 담당자를 통해 진행합니다."],
        ],
        [4.5, 4.5, 4.5, 4.5],
    )

    add_heading(doc, "13. FAQ / 자주 묻는 질문")
    add_table(
        doc,
        ["질문", "답변"],
        [
            [
                "재작업을 하면 기존 결과가 지워지나요?",
                "기존 완료 이력은 그대로 남고, 재작업으로 만들어진 새 단계가 추가됩니다.",
            ],
            [
                "어떤 옵션을 선택해야 할지 잘 모르겠습니다.",
                "현재 단계만 다시 하면 충분하면 \"현재 활동만\", 이후 단계 결과도 다시 만들어야 하면 \"이후 모든 활동 포함\"을 선택합니다.",
            ],
            [
                "재작업 시작 버튼을 눌렀는데 다시 이전 단계가 \"대기\"로 보이는 것이 정상인가요?",
                "정상입니다. 시작 단계가 \"진행 중\"으로 표시되고, 함께 선택된 이후 단계는 \"대기\" 상태로 줄지어 표시됩니다.",
            ],
            [
                "잘못된 범위로 재작업을 시작했습니다. 어떻게 하나요?",
                "운영 담당자에게 인스턴스 ID와 함께 상황을 공유하고 후속 조치를 협의합니다.",
            ],
            [
                "재작업 버튼이 다른 사용자의 작업에서는 왜 보이지 않나요?",
                "재작업은 본인이 완료한 작업에서만 시작할 수 있습니다.",
            ],
        ],
        [6.0, 11.0],
    )

    add_heading(doc, "14. 효과적인 재작업 사용 팁")
    add_bullets(
        doc,
        [
            "재작업 전에 어떤 결과를 다시 만들고 싶은지 한 문장으로 정리합니다. 예: \"승인 단계의 의견을 수정하고 이후 통보까지 다시 보낸다.\"",
            "처음에는 \"현재 활동만\" 옵션으로 영향 범위를 최소화한 뒤, 필요할 때 \"이후 모든 활동 포함\" 옵션으로 확장합니다.",
            "다이얼로그에서 강조된 활동 칩이 실제로 다시 수행되어야 할 단계와 일치하는지 한 번 더 확인한 후 \"재작업 시작\" 버튼을 누릅니다.",
            "재작업 후 인스턴스 상세 화면에서 \"진행 중\" 단계가 의도한 단계와 일치하는지 확인하는 습관을 들입니다.",
            "동일 인스턴스에서 짧은 간격으로 여러 번 재작업하는 대신, 필요한 범위를 한 번에 선택해 시작하면 단계 상태를 추적하기 쉽습니다.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
