"""Generate the end-user manual DOCX for completion_process-workitem-submission.

Based on the bundled STYLE_REFERENCE.py from the docx-user-manual skill.
Run: python generate_completion_process-workitem-submission_user_manual.py
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SPEC_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = SPEC_ROOT / "docs" / "completion_process-workitem-submission-user-manual.docx"
SCREENSHOT_DIR = SPEC_ROOT / "e2e" / "results" / "screenshots"

DOC_VERSION = "v1.0"
DOC_DATE = date(2026, 5, 27)
ORG_NAME = "Process-GPT"
SERVICE_NAME = "Process-GPT 워크아이템"
DOC_TITLE = "프로세스 워크아이템 처리 사용자 매뉴얼"
FONT_NAME = "Malgun Gothic"
PRIMARY_COLOR = "1F4E79"
CAPTION_COLOR = "666666"
TABLE_HEADER_FILL = "D9EAF7"


SCREENSHOTS = {
    "todolist_form": "process-gpt-completion_process-workitem-submission-04-todolist-form.png",
    "todolist_submitted": "process-gpt-completion_process-workitem-submission-04-todolist-submitted.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차는 Word에서 필드 업데이트 시 페이지 번호와 함께 표시됩니다."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def enable_field_update_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, bold=True, color=PRIMARY_COLOR)


def add_paragraph(doc: Document, text: str = "", style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.45)
        run = paragraph.add_run(f"{idx}.\t{item}")
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, TABLE_HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_image(doc: Document, key: str, caption: str) -> None:
    path = SCREENSHOT_DIR / SCREENSHOTS[key]
    if not path.exists():
        add_paragraph(doc, f"[화면 예시 누락] {path.name}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=CAPTION_COLOR)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SERVICE_NAME)
    set_run_font(run, size=18, bold=True, color=PRIMARY_COLOR)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DOC_TITLE)
    set_run_font(run, size=28, bold=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["문서 버전", DOC_VERSION],
            ["작성일", DOC_DATE.isoformat()],
            ["작성자/조직", ORG_NAME],
            ["문서 구분", "사용자 참조 매뉴얼"],
            ["대상 서비스", SERVICE_NAME],
        ],
        [4.0, 12.0],
    )
    doc.add_page_break()


def add_scenario(
    doc: Document,
    title: str,
    purpose: str,
    preconditions: list[str],
    steps: list[str],
    expected: list[str],
    images: list[tuple[str, str]] | None = None,
) -> None:
    add_heading(doc, title, level=2)
    add_paragraph(doc, f"이럴 때 사용합니다: {purpose}")
    add_paragraph(doc, "사용 전 확인")
    add_bullets(doc, preconditions)
    add_paragraph(doc, "따라 하기")
    add_numbers(doc, steps)
    add_paragraph(doc, "화면에서 확인할 내용")
    add_bullets(doc, expected)
    if images:
        for key, caption in images:
            add_image(doc, key, caption)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{SERVICE_NAME} 사용자 매뉴얼 | ")
    set_run_font(run, size=9, color=CAPTION_COLOR)
    add_page_number(footer)

    doc.core_properties.title = DOC_TITLE
    doc.core_properties.subject = SERVICE_NAME
    doc.core_properties.author = ORG_NAME
    enable_field_update_on_open(doc)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    # 1. 문서 정보 및 변경 이력
    add_heading(doc, "1. 문서 정보 및 변경 이력")
    add_table(
        doc,
        ["문서 버전", "변경일", "변경 내용", "작성자", "검토자"],
        [[DOC_VERSION, DOC_DATE.isoformat(), "초안 작성", ORG_NAME, "서비스 담당자"]],
        [2.6, 3.0, 8.2, 3.0, 3.0],
    )

    # 2. 목차
    add_heading(doc, "2. 목차")
    add_paragraph(doc, "아래 목차는 Word에서 문서를 열 때 필드를 업데이트하면 페이지 번호가 반영됩니다.")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    # 3. 개요
    add_heading(doc, "3. 개요")
    add_paragraph(
        doc,
        "이 문서는 Process-GPT 워크아이템 화면을 처음 사용하는 분이 할 일 목록에서 본인에게 배정된 작업을 열고 "
        "필요한 값을 입력한 뒤 제출까지 완료할 수 있도록 안내합니다.",
    )
    add_bullets(
        doc,
        [
            "서비스 목적: 진행 중인 업무 절차의 각 단계를 사용자가 화면에서 처리하고, 처리 결과를 다음 단계로 넘길 수 있도록 돕습니다.",
            "매뉴얼 목적: 로그인, 작업 열기, 입력값 작성, 제출, 결과 확인까지 사용자가 따라 할 수 있도록 안내합니다.",
            "기능 범위: 할 일 화면 진입, 워크아이템 폼 입력, 제출, 제출 결과 확인을 포함합니다.",
        ],
    )

    # 4. 대상 사용자
    add_heading(doc, "4. 대상 사용자")
    add_table(
        doc,
        ["사용자 역할", "주요 사용 목적", "권장 사용 범위"],
        [
            ["담당자", "본인에게 배정된 워크아이템 폼을 열고 값을 입력하여 제출합니다.", "할 일 처리 흐름 전체"],
            ["검토자", "제출된 워크아이템 결과를 확인하고 다음 단계의 입력 자료로 활용합니다.", "결과 확인 및 후속 검토"],
        ],
        [3.5, 8.5, 5.0],
    )

    # 5. 사용 전 확인 사항
    add_heading(doc, "5. 사용 전 확인 사항")
    add_bullets(
        doc,
        [
            "Process-GPT 계정(이메일과 비밀번호)이 발급되어 있어야 합니다.",
            "본인에게 할당된 워크아이템(할 일)이 한 건 이상 존재해야 합니다. 새 작업은 담당자에게 자동으로 배정됩니다.",
            "최신 버전의 크롬, 엣지, 사파리 등 표준 웹 브라우저에서 사용을 권장합니다.",
            "서비스 접속 URL과 접속 가능한 사내망/인터넷 환경이 준비되어 있어야 합니다.",
            "처리할 작업에 입력해야 하는 정보(예: 사용 일수, 사유 등)를 미리 확인해 두면 편리합니다.",
        ],
    )

    # 6. 시작하기
    add_heading(doc, "6. 시작하기")
    add_paragraph(doc, "처음 접속해서 본인 할 일을 열기까지의 기본 순서를 안내합니다.")
    add_numbers(
        doc,
        [
            "브라우저 주소창에 서비스 접속 URL을 입력합니다.",
            "로그인 화면에서 이메일과 비밀번호를 입력하고 로그인 버튼을 누릅니다.",
            "로그인이 완료되면 본인에게 배정된 할 일 목록 화면으로 이동합니다.",
            "처리할 항목을 선택해 해당 워크아이템의 작업 화면을 엽니다.",
            "화면에 표시된 입력 항목을 채운 뒤 제출 버튼으로 처리를 완료합니다.",
        ],
    )

    # 7. 화면 구성
    add_heading(doc, "7. 화면 구성")
    add_paragraph(doc, "워크아이템 작업 화면은 좌측의 보조 영역과 우측의 작업 본문 영역으로 구성됩니다.")
    add_table(
        doc,
        ["화면 영역", "역할"],
        [
            ["좌측 보조 영역", "탐색 메뉴와 진행 중인 작업의 보조 정보를 제공합니다."],
            ["상단 정보 영역", "현재 열려 있는 워크아이템의 제목과 상태를 표시합니다."],
            ["본문 작업 영역", "사용자가 값을 입력하거나 확인할 수 있는 폼이 표시됩니다."],
            ["하단 동작 영역", "작성한 내용을 저장하거나 다음 단계로 제출하는 버튼이 위치합니다."],
        ],
        [5.0, 11.0],
    )
    add_image(
        doc,
        "todolist_form",
        "그림 1. 할 일 목록에서 본인에게 배정된 워크아이템 작업 화면을 연 첫 화면",
    )

    # 8. 주요 사용 흐름
    add_heading(doc, "8. 주요 사용 흐름")
    add_scenario(
        doc,
        "8.1 본인 할 일 워크아이템 제출하기",
        "본인에게 배정된 할 일을 열어 필요한 값을 입력한 뒤 제출 처리하여 다음 단계로 넘기는 가장 일반적인 사용 흐름입니다.",
        [
            "Process-GPT 계정(이메일과 비밀번호)이 준비되어 있습니다.",
            "본인 할 일 목록에 처리 대상 워크아이템이 한 건 이상 표시됩니다.",
            "워크아이템에 입력해야 할 값(예: 사용 일수 등)을 미리 확인했습니다.",
        ],
        [
            "브라우저로 서비스 로그인 화면에 접속합니다.",
            "이메일과 비밀번호를 입력하고 로그인합니다.",
            "할 일 목록에서 처리할 항목을 선택해 워크아이템 작업 화면을 엽니다.",
            "본문 영역의 입력 항목에 필요한 값을 입력합니다(예: 사용 일수에 숫자 입력).",
            "하단의 제출 버튼을 눌러 워크아이템 제출을 완료합니다.",
        ],
        [
            "제출이 정상 처리되면 화면이 제출 완료 상태로 갱신됩니다.",
            "입력한 값이 화면에 그대로 표시되어 내가 보낸 값이 반영되었음을 확인할 수 있습니다.",
            "제출된 작업은 더 이상 본인이 새로 입력해야 하는 할 일 목록에서 진행 대상이 아니게 됩니다.",
        ],
        images=[
            ("todolist_form", "그림 2. 워크아이템 폼을 연 직후 입력값을 작성하기 전의 화면"),
            ("todolist_submitted", "그림 3. 제출이 완료된 후 결과 상태가 확정된 화면"),
        ],
    )

    add_scenario(
        doc,
        "8.2 이미 시작된 작업을 이어서 처리하기",
        "이전에 누군가가 시작해 둔 인스턴스나 일정에 따라 자동으로 생성된 작업을 본인 계정으로 이어받아 마무리하는 흐름입니다.",
        [
            "본인에게 새로 배정된 할 일이 목록에 표시되어 있습니다.",
            "해당 작업이 어떤 업무 단계인지 사전에 안내받았거나 작업 제목으로 확인할 수 있습니다.",
        ],
        [
            "할 일 목록에서 이어서 처리할 워크아이템을 엽니다.",
            "이미 채워져 있는 값이 있다면 내용을 확인합니다.",
            "추가/수정이 필요한 항목을 입력합니다.",
            "제출 버튼을 눌러 작업을 마무리합니다.",
        ],
        [
            "기존 작업이 새로운 입력값으로 갱신된 채 제출 완료 상태로 바뀝니다.",
            "동일한 작업을 다시 열어도 제출된 결과를 확인할 수 있습니다.",
        ],
    )

    # 9. 화면 항목 및 옵션 설명
    add_heading(doc, "9. 화면 항목 및 옵션 설명")
    add_paragraph(doc, "워크아이템 작업 화면에서 자주 보이는 항목과 그 의미를 설명합니다. 실제 항목 이름과 기본값은 업무 흐름에 따라 다를 수 있습니다.")
    add_table(
        doc,
        ["화면 항목/옵션", "사용 여부", "설명"],
        [
            ["작업 제목", "자동 표시", "현재 처리하는 업무 단계의 이름을 보여 줍니다."],
            ["입력 항목(예: 사용 일수)", "필수", "업무 단계에서 요구하는 값을 입력합니다. 숫자/날짜 등 항목별 형식에 맞춰 입력합니다."],
            ["보조 입력 항목", "선택", "필요할 때 추가 정보를 함께 입력할 수 있는 영역입니다."],
            ["제출 버튼", "필수", "작성한 값을 다음 단계로 넘기고 워크아이템을 제출 완료 상태로 전이합니다."],
            ["취소/이전", "선택", "작업 화면에서 나가거나 이전 화면으로 돌아갑니다."],
        ],
        [4.5, 3.0, 9.5],
    )

    # 10. 결과 확인
    add_heading(doc, "10. 결과 확인")
    add_bullets(
        doc,
        [
            "제출이 정상 처리되면 화면이 제출 완료 상태로 갱신됩니다.",
            "내가 입력한 값이 화면에 그대로 표시되어 어떤 내용이 다음 단계로 전달되었는지 확인할 수 있습니다.",
            "제출이 완료된 워크아이템은 본인이 다시 입력해야 하는 진행 대상에서 제외됩니다.",
            "필요 시 동일한 워크아이템을 다시 열어 제출된 결과를 확인할 수 있습니다.",
        ],
    )

    # 11. 오류 및 예외 상황
    add_heading(doc, "11. 오류 및 예외 상황")
    add_table(
        doc,
        ["증상", "주요 원인", "해결 방법"],
        [
            [
                "로그인이 되지 않습니다.",
                "이메일/비밀번호가 잘못되었거나 계정이 아직 발급되지 않았습니다.",
                "이메일과 비밀번호를 다시 확인하고, 계속 실패하면 서비스 담당자에게 계정 발급 또는 비밀번호 초기화를 요청합니다.",
            ],
            [
                "할 일 목록에 작업이 보이지 않습니다.",
                "본인에게 배정된 작업이 아직 없거나, 잘못된 계정으로 로그인했을 수 있습니다.",
                "올바른 계정으로 로그인되었는지 확인하고, 업무 담당자에게 본인 계정으로 작업이 배정되었는지 문의합니다.",
            ],
            [
                "워크아이템 화면에 폼이 표시되지 않습니다.",
                "해당 워크아이템이 이미 처리되었거나 본인에게 배정되지 않은 작업일 수 있습니다.",
                "할 일 목록으로 돌아가 다른 항목을 선택하거나, 업무 담당자에게 해당 작업의 현재 상태를 확인합니다.",
            ],
            [
                "제출 버튼을 눌렀는데 결과가 갱신되지 않습니다.",
                "네트워크가 일시적으로 끊겼거나 필수 항목이 비어 있을 수 있습니다.",
                "필수 입력 항목을 채웠는지 확인한 뒤 다시 제출하고, 계속 실패하면 화면을 새로고침해 본인 할 일 목록에서 다시 시도합니다.",
            ],
            [
                "이미 제출한 작업이 다시 진행 대상으로 보입니다.",
                "다른 단계에서 새로운 워크아이템이 생성되었을 수 있습니다.",
                "작업 제목과 단계를 확인해 새로 배정된 작업인지 판단한 뒤 다시 처리합니다.",
            ],
        ],
        [4.0, 5.5, 7.5],
    )

    # 12. 권한 및 역할별 기능 차이
    add_heading(doc, "12. 권한 및 역할별 기능 차이")
    add_table(
        doc,
        ["역할", "사용 가능한 기능", "제한 사항"],
        [
            [
                "담당자",
                "본인 할 일 목록 확인, 워크아이템 작업 화면 진입, 폼 입력 및 제출.",
                "본인에게 배정되지 않은 워크아이템은 처리할 수 없습니다.",
            ],
            [
                "검토자",
                "제출된 워크아이템 결과 확인 및 후속 검토 단계 처리.",
                "다른 담당자에게 배정된 입력 단계는 직접 처리할 수 없습니다(권한 필요).",
            ],
            [
                "운영/관리자",
                "프로세스 정의 등록, 인스턴스 운영 관리 등의 관리 기능.",
                "일반 사용자 매뉴얼 범위 외이므로 별도 운영 가이드를 참고합니다(권한 필요).",
            ],
        ],
        [3.5, 7.0, 6.5],
    )

    # 13. FAQ
    add_heading(doc, "13. FAQ / 자주 묻는 질문")
    add_table(
        doc,
        ["질문", "답변"],
        [
            [
                "할 일 목록에 작업이 새로 생기지 않습니다. 어떻게 해야 하나요?",
                "본인 계정으로 로그인되었는지 먼저 확인하고, 업무 담당자에게 본인 앞으로 작업이 배정되었는지 문의합니다.",
            ],
            [
                "잘못된 값을 입력해 제출했습니다. 수정할 수 있나요?",
                "제출 완료된 작업은 일반적으로 본인 화면에서 다시 수정할 수 없습니다. 후속 단계 담당자나 운영자에게 정정 요청을 전달합니다.",
            ],
            [
                "한 작업을 여러 사람이 처리할 수 있나요?",
                "기본적으로 워크아이템은 배정된 담당자가 처리합니다. 역할에 따라 배정 규칙이 다를 수 있으니 업무 담당자에게 문의합니다.",
            ],
            [
                "제출 후 다음 단계는 누가 처리하나요?",
                "제출이 완료되면 프로세스 정의에 따라 다음 단계 담당자에게 자동으로 작업이 전달됩니다.",
            ],
            [
                "테넌트(소속 조직)가 다른 작업도 보이나요?",
                "본인이 로그인한 조직(테넌트)의 작업만 표시되며, 다른 조직의 정보는 보이지 않습니다.",
            ],
        ],
        [6.5, 11.5],
    )

    # 14. 효과적인 사용 팁
    add_heading(doc, "14. 효과적인 사용 팁")
    add_bullets(
        doc,
        [
            "작업을 열기 전 어떤 값을 입력해야 하는지 미리 확인하면 작성 시간이 짧아집니다.",
            "여러 항목을 입력해야 할 때는 위에서 아래로 한 번에 작성한 뒤 제출하는 것이 안전합니다.",
            "오류 메시지가 나오면 화면을 닫지 말고 메시지에 표시된 안내(필수 항목, 형식)에 따라 값을 다시 확인합니다.",
            "잘못 제출한 경우 임의로 다시 시도하지 말고, 업무 담당자에게 정정 절차를 먼저 문의합니다.",
            "작업이 갑자기 사라졌다면 다른 단계로 넘어갔을 가능성이 있으니 작업 제목으로 진행 단계를 확인합니다.",
        ],
    )

    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
